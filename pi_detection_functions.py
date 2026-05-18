import io
import os
import re
import shutil
import tempfile
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List

import openpyxl
import pandas as pd
import pdfplumber
import pytesseract
import spacy
from PIL import Image, ImageOps, ImageFilter
from PyPDF2 import PdfReader
from docx import Document
from zipfile import ZipFile
import xml.etree.ElementTree as ET
import posixpath

# Optional PDF page rendering for OCR fallback
try:
    import pypdfium2 as pdfium
except Exception:
    pdfium = None  # page render OCR will be skipped if unavailable

# Only set tesseract_cmd if the binary exists at this path (Windows users)
_win_tesseract = r"C:\Users\BJ574SU\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"
if os.path.exists(_win_tesseract):
    pytesseract.pytesseract.tesseract_cmd = _win_tesseract


# Load spaCy once and cache it
from functools import lru_cache

import streamlit as st
@st.cache_resource(show_spinner=False)
def get_nlp():
    """Load spaCy's en_core_web_sm once and reuse it. Returns None if the model isn't available."""
    try:
        import spacy  # local import avoids import errors at module import time
    except Exception:
        return None

    try:
        nlp = spacy.load(
            "en_core_web_sm",
            disable=["tagger", "parser", "attribute_ruler", "lemmatizer"]
        )
        nlp.max_length = max(getattr(nlp, "max_length", 1_000_000), 2_000_000)
        return nlp
    except Exception:
        import os
        local_dir = os.environ.get("SPACY_LOCAL_MODEL_DIR")
        if local_dir:
            try:
                nlp = spacy.load(
                    local_dir,
                    disable=["tagger", "parser", "attribute_ruler", "lemmatizer"]
                )
                nlp.max_length = max(getattr(nlp, "max_length", 1_000_000), 2_000_000)
                return nlp
            except Exception:
                pass
        return None

# REGEX PATTERNS
pi_PATTERNS: Dict[str, str] = {
    # Common PII
    "email": r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b",

    # Phone requires formatting chars to avoid matching bare digit strings
    "Phone": r"(?<!\d)(?:\+\d{1,3}[\s\-]?)?\(?\d{3}\)?[\s\-\.]\d{3}[\s\-\.]\d{4}\b",

    # Account Number is context-gated in get_confidence (requires banking keywords)
    "Account Number": r"\b\d{9,18}\b",

    "SSN Number": r"\b\d{3}[\s]?-\d{2}[\s]?-\d{4}\b",

    # Passports
    "US Passport Number": r"(?<![A-Z\-])\b(?:[A-Z]\d{8}|\d{9})\b",

    # Age mentions
    "Age": r"\b(?:1[01]\d|[1-9]?\d|120)\s*(?:years?|yrs?)\b",
    "Gender": r"\b(?:(?:male|female|transgender|non[-\s]?binary)|m/f)\b",
    "Credit/Debit Card": r"\b(?:3[47]\d{2}[-\s]?\d{6}[-\s]?\d{5}|(?:\d{4}[-\s]?){3}\d{4})\b",

    # Indian PII
    # Aadhaar: 12 digits grouped XXXX XXXX XXXX, first digit 2-9
    "Aadhaar Number": r"\b[2-9]\d{3}[\s\-]{0,2}\d{4}[\s\-]{0,2}\d{4}\b",
    # PAN Card: 5 letters + 4 digits + 1 letter e.g. ABCDE1234F
    "PAN Card":        r"\b[A-Z]{5}[0-9]{4}[A-Z]\b",
    # IFSC: 4-5 letters + 0 + 6 alphanumeric e.g. HDFC0001234, KOTAK0005678
    "IFSC Code":       r"\b[A-Z]{4,5}0[A-Z0-9]{6}\b",
    # Indian phone: starts 6-9, exactly 10 digits, optional +91 prefix
    "Indian Phone":    r"\b(?:\+91[\s\-]?)?[6-9]\d{9}\b",
    # Voter ID: 3 uppercase letters + 7 digits
    "Voter ID":        r"\b(?!MRN)[A-Z]{2,3}[0-9]{7}\b",
    # IP Address
    "IP Address":      r"\b(?:(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\.){3}(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\b",
}

_ACCOUNT_KW_RE = re.compile(
    r"\b(?:account|acct|a\/c|bank|iban|swift|ach|sort\s*code|"
    r"wire\s*transfer|wire|transfer|direct\s*deposit|checking|savings|"
    r"beneficiary|remit|remittance|payee|debit|credit)\b",
    re.IGNORECASE,
)
# Routing number keyword (9-digit ABA routing numbers)
_ROUTING_KW_RE = re.compile(r"\b(?:routing|aba|sort\s*code)\b", re.IGNORECASE)
_ROUTING_NUM_RE = re.compile(r"\b\d{9}\b")  # US ABA routing = exactly 9 digits

# Checks if an explicit account keyword appears directly before the number (within 25 chars)
# Used to override the Indian-phone heuristic for 10-digit numbers starting 6-9
_EXPLICIT_ACCT_KW_RE = re.compile(
    r"\b(?:account|acct|a\/c|bank\s+account|account\s+(?:no|number|num))\s*[:\|]?\s*$",
    re.IGNORECASE,
)

def _has_account_context(ctx: str, col_name: str = "") -> bool:
    return bool(
        _ACCOUNT_KW_RE.search(ctx or "") or
        _ROUTING_KW_RE.search(ctx or "") or
        "account" in (col_name or "").lower() or
        bool(re.search(r"\b(?:ACH|IBAN|SWIFT|IFSC)\b", ctx or ""))
    )

def _is_explicit_account(value: str, full_text: str) -> bool:
    """Return True if an explicit account keyword appears directly before this value (within 25 chars)."""
    idx = full_text.find(value)
    if idx < 0:
        return False
    window = full_text[max(0, idx - 25):idx]
    return bool(_EXPLICIT_ACCT_KW_RE.search(window))

def _is_routing_number(value: str, context: str) -> bool:
    """Return True only if a routing keyword appears within 30 chars before this value."""
    digits = re.sub(r"\D", "", value)
    if len(digits) != 9:
        return False
    if not _ROUTING_KW_RE.search(context or ""):
        return False
    # Only check the text before the value — not after (prevents passport being tagged)
    idx = (context or "").find(value)
    if idx < 0:
        return False
    window_before = context[max(0, idx - 30): idx]
    return bool(_ROUTING_KW_RE.search(window_before))

_VOTER_ID_CTX_RE = re.compile(r"\b(?:voter|epic|election|electoral)\b", re.IGNORECASE)
_DL_CTX_STRONG_RE = re.compile(r"\b(?:driver\'?s?|driving|license|licence|dl)\b", re.IGNORECASE)
def _resolve_voter_id_vs_dl(ctx: str) -> str:
    if _VOTER_ID_CTX_RE.search(ctx): return "Voter ID"
    if _DL_CTX_STRONG_RE.search(ctx): return "Driving License"
    return "Voter ID"

_MRN_CTX_RE = re.compile(
    r"\b(?:mrn|medical\s*record(?:\s*(?:no|number|#))?|patient\s*(?:id|no))\b",
    re.IGNORECASE,
)
def _has_mrn_context(ctx: str) -> bool:
    return bool(_MRN_CTX_RE.search(ctx or ""))

# US Passport
PASSPORT_KEYWORD = re.compile(r"\bpassport\b", re.IGNORECASE)
PASSPORT_STRICT_RE = re.compile(r"^[A-Z]\d{8}$")
PASSPORT_9D_RE     = re.compile(r"^\d{9}$")

def is_passport_token(v: str, ctx: str) -> bool:
  v = (v or "").strip()
  if PASSPORT_STRICT_RE.match(v):
      return True
  return bool(PASSPORT_9D_RE.match(v) and PASSPORT_KEYWORD.search(ctx or ""))

ADDRESS_PATTERN = (
    r"\b\d{1,6}\s+"
    r"(?:[A-Za-z0-9.\-]{2,}\s+){0,6}"
    r"(?:Street|St\.?|Avenue|Ave\.?|Boulevard|Blvd\.?|Road|Rd\.?|Lane|Ln\.?|"
    r"Drive|Dr\.?|Court|Ct\.?|Circle|Cir\.?|Highway|Hwy\.?|Place|Pl\.?|"
    r"Terrace|Ter\.?|Parkway|Pkwy\.?|Broadway|Way|Nagar|Marg|Chowk|"
    r"Salai|Lines|Hills|Lake|Colony|Layout|Extension|Cross|Sector|Block|Phase|Vihar|Enclave)"
    r"(?:[,\s]+(?:Apt|Apartment|Unit|Suite|Ste|Flat|Floor|#)\s*[\w\-]+)?"
    r"(?:[,\s]+[A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){0,2})?"
    r"(?:[,\s]+[A-Z]{2})?"
    r"(?:[,\s]+\d{5,6})?"
)

ADDRESS_SUITE_PATTERN = re.compile(
    r"\b(?:Suite|Ste|Unit|Apt|Apartment|Flat|Floor|#)\s*[\w\-]+[,\s]+"
    r"[A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){0,3}"   # city
    r"(?:[,\s]+[A-Z]{2})?"                           # state abbrev
    r"(?:[,\s]+\d{5,6})?",                           # zip
    re.IGNORECASE
)

# US Driving Licenses
US_DL_PATTERNS: Dict[str, str] = {
    "AL": r"\b\d{8}\b",  # Exact 8 digits
    "AK": r"\b\d{7}\b",
    "AZ": r"\b(?:AZ[\s\-–—]*)?(?P<az_core>[A-Z]\d{8})\b",
    "AR": r"\b\d{9}\b",  # Exact 9 digits (was 4,9)
    "CA": r"\b(?:CA[\s\-–—]*)?(?P<core>[A-Z]\d{7})\b", 
        "CO": r"\b(?:CO[\s\-–—]*)?(?P<co_core>[A-Z]\d{7,8})\b",  # letter+digits only; bare \d{9} removed to avoid routing/passport confusion
    "CT": r"\b\d{9}\b",
    "DE": r"\b\d{7}\b",  # Exact 7 digits (was 1,7)
    "DC": r"\b\d{7}\b",
    "FL": r"\b(?:FL[\s\-–—]*)?[A-Z]\d{12}\b",
    "GA": r"\b\d{9}\b",
    "HI": r"\b[A-Z]\d{9}\b",
    "ID": r"\b[A-Z]{2}\d{6}\b",
    "IL": r"\b(?:IL[\s\-–—]*)?[A-Z]\d{11}\b",
    "IN": r"\b\d{10}\b",
    "IA": r"\b\d{3}[- ]?\d{2}[- ]?\d{4}\b",
    "KS": r"\b(?:KS[\s\-–—]*)?[A-Z]\d{8}\b",
    "KY": r"\b(?:KY[\s\-–—]*)?[A-Z]\d{9}\b",
    "LA": r"\b\d{9}\b",
    "ME": r"\b\d{7}\b",
    "MD": r"\b[A-Z]{3}\d{9}\b",  # Exact 3 letters (was 1,3)
    "MA": r"\b(?:MA[\s\-–—]*)?[A-Z]\d{8}\b",
    "MI": r"\b(?:MI[\s\-–—]*)?[A-Z]\d{12}\b",
    "MN": r"\b(?:MN[\s\-–—]*)?[A-Z]\d{12}\b",
    "MS": r"\b\d{9}\b",
    "MO": r"\b(?:MO[\s\-–—]*)?[A-Z]\d{9}\b",
    "MT": r"\b(?:MT[\s\-–—]*)?[A-Z]\d{8}\b",
    "NE": r"\b(?:NE[\s\-–—]*)?[A-Z]\d{8}\b",
    "NV": r"\b\d{10}\b",
    "NH": r"\b\d{2}[A-Z]\d{8}\b",
    "NJ": r"\b\d{2}[A-Z]\d{6}\b",
    "NM": r"\b\d{9}\b",  # Exact 9 digits (was 8,9)
    "NY": r"\b(?:NY[\s\-–—]*)?(?P<ny_core>[A-Z]\d{8})\b",
    "NC": r"\b\d{12}\b",  # Exact 12 digits (was 1,12)
    "ND": r"\b\d{13}\b",
    "OH": r"\b(?:OH[\s\-–—]*)?(?P<oh_core>[A-Z]\d{8})\b",  # Exact 8 digits after letter (was 4,8)
    "OK": r"\b(?:OK[\s\-–—]*)?[A-Z]\d{9}\b",
    "OR": r"\b\d{8}\b",
    "PA": r"\b\d{8}\b",
    "RI": r"\b\d{7}\b",
    "SC": r"\b\d{11}\b",  # Exact 11 digits (was 5,11)
    "SD": r"\b\d{12}\b",  # Exact 12 digits (was 8,12)
    "TN": r"\b\d{9}\b",  # Exact 9 digits (was 7,9)
    "TX": r"\b(?:TX[\s\-–—]*)?(?P<core>[A-Z]\d{8}|\d{8})\b",
    "UT": r"\b\d{10}\b",  # Exact 10 digits (was 4,16)
    "VT": r"\b\d{8}\b",
    "VA": r"\b[A-Z]\d{11}\b",  # Exact 11 digits after letter (was 8,11)
    "WA": r"\b[A-Z]{2}\d{12}\b",
    "WV": r"\b[A-Z]\d{6}\b",
    "WI": r"\b[A-Z]\d{13}\b",
    "WY": r"\b\d{9}\b",
}

# Iowa DL pattern removed to avoid SSN collision (matches 3-2-4 like SSN)
US_DL_PATTERNS.pop("IA", None)

# States whose DL pattern is pure digits — very ambiguous.
# We only emit them when a DL-specific column header is present (not just DL keyword in text).
_PURE_DIGIT_DL_STATES = {
    "AR", "GA", "IN", "LA", "ME", "MS", "NM", "NV", "OR", "PA",
    "RI", "SC", "SD", "TN", "UT", "VT", "WY", "NC", "ND", "CT", "DC"
}

US_DL_MAX_LENGTH: Dict[str, int] = {
    "AL": 8, "AK": 7, "AZ": 9, "AR": 9, "CA": 8, "CO": 9, "CT": 9, "DC": 7, "DE": 7, "FL": 13,
    "GA": 9, "HI": 10, "ID": 8, "IL": 12, "IN": 10, "IA": 9, "KS": 9, "KY": 10, "LA": 9,
    "ME": 7, "MD": 12, "MA": 9, "MI": 13, "MN": 13, "MS": 9, "MO": 10, "MT": 9, "NE": 9,
    "NV": 10, "NH": 11, "NJ": 9, "NM": 9, "NY": 9, "NC": 12, "ND": 13, "OH": 9, "OK": 10,
    "OR": 8, "PA": 8, "RI": 7, "SC": 11, "SD": 12, "TN": 9, "TX": 9, "UT": 10, "VT": 8,
    "VA": 12, "WA": 14, "WV": 7, "WI": 14, "WY": 9
}

# Minimum DL length for OCR validation
US_DL_MIN_LENGTH: Dict[str, int] = {
    "AL": 6, "AK": 7, "AZ": 9, "AR": 4, "CA": 8, "CO": 8, "CT": 9, "DE": 1, "FL": 13,
    "GA": 7, "HI": 9, "ID": 8, "IL": 12, "IN": 10, "IA": 8, "KS": 9, "KY": 10, "LA": 9,
    "ME": 7, "MD": 1, "MA": 9, "MI": 13, "MN": 13, "MS": 9, "MO": 10, "MT": 9, "NE": 9,
    "NV": 10, "NH": 10, "NJ": 8, "NM": 8, "NY": 9, "NC": 1, "ND": 13, "OH": 1, "OK": 10,
    "OR": 8, "PA": 8, "RI": 7, "SC": 5, "SD": 8, "TN": 7, "TX": 8, "UT": 4, "VT": 8,
    "VA": 9, "WA": 14, "WV": 7, "WI": 14, "WY": 9
}

# Company IDs (both numeric 6-digit and alphanumeric EMP-prefixed)
COMPANY_ID_PATTERN_NUM = r"\b\d{6}\b"
COMPANY_ID_RE_NUM = re.compile(COMPANY_ID_PATTERN_NUM)

COMPANY_ID_CONTEXT_RE = re.compile(
    r"\b(?:company|employee|emp|staff|badge|payroll)\s*(?:id|#|number|no\.?|code|num)\b"
    r"|\bcompany\s*id\s*[:\|]"   # "Company ID:" label
    r"|\bemp(?:loyee)?\s*[:\|]", # "EMP:" or "Employee:" label
    re.IGNORECASE
)

# Also detect EMP-prefixed company IDs directly (e.g. EMP4821, EMP7721)
ALNUM_COMPANY_ID_RE = re.compile(r"\b(?:EMP\d{4,8}|[A-Z]{1,3}\d{4,8})\b")

COMPANY_ID_NEG_CONTEXT_RE = re.compile(
    r"\b(?:mrn|medical\s*record|insurance\s*id|policy(?:\s*(?:no|number))?|passport|routing|account|ach|iban|visa|mastercard|voter|aadhaar|aadhar|pan\s*card|dl|driver)\b",
    re.IGNORECASE
)

def is_company_id(token: str, context: str | None = None) -> bool:
    if not token:
        return False
    looks_like = bool(ALNUM_COMPANY_ID_RE.fullmatch(token) or COMPANY_ID_RE_NUM.fullmatch(token))
    if not looks_like:
        return False
    # EMP-prefixed values are unambiguous
    if re.match(r'^EMP\d+$', token, re.IGNORECASE):
        return True
    # reject MRN-prefixed
    if re.match(r'^MRN', token, re.IGNORECASE):
        return False
    # reject passport format
    if re.match(r'^[A-Z]\d{8}$', token):
        return False
    # reject Voter ID format
    if re.match(r'^[A-Z]{2,3}\d{7}$', token):
        return False
    # reject INS-prefixed insurance IDs
    if re.match(r'^INS', token, re.IGNORECASE):
        return False
    ctx = context or ""
    if COMPANY_ID_NEG_CONTEXT_RE.search(ctx):
        return False
    return bool(COMPANY_ID_CONTEXT_RE.search(ctx))

# MRN / Insurance ID
MRN_RE = re.compile(
    r"\bMRN[\s:#-]*([A-Z0-9][A-Z0-9\-]{3,12})\b"   # MRN-KIM-4455, MRN0456789
    r"|\bMRN[\s:#-]*([A-Z]{0,3}\d{6,10})\b",         # MRN followed by digits (keyword required)
    re.IGNORECASE
)
INSURANCE_ID_RE = re.compile(
    r"\b(?:insurance\s*id|policy\s*id|policy\s*no\.?)\s*[:#-]?\s*([A-Z0-9][A-Z0-9\-]{5,20})\b"
    r"|\b(INS-[A-Z0-9][A-Z0-9\-]{4,18})\b",  # direct INS- prefix match
    re.IGNORECASE
)

# SSN with OCR tolerance (handles spacing/dash variants)
SSN_NEARBY_FLEX_RE = re.compile(
    r"(?:(?:ssn|social\s*security)\s*[:#-]?\s*)(\d{3})[-\s]?(\d{2})[-\s]?(\d{4})\b",
    re.IGNORECASE
)

# Excel column header hints that indicate the column is an ID, not a DL
ID_COL_HINT_RE = re.compile(
    r"\b(emp(?:loyee)?\s*(?:id|code)|company\s*id|staff\s*id|payroll\s*id|badge\s*id|id)\b",
    re.IGNORECASE
)

# Excel column header hints for name columns, used for standalone person name detection
NAME_COL_HINT_RE = re.compile(
    r"(?<![A-Za-z])(?:employee[\s_]?name|emp[\s_]?name|full[\s_]?name|fullname"
    r"|first[\s_]?name|firstname|last[\s_]?name|lastname|name)(?![A-Za-z])",
    re.IGNORECASE
)

# Matches full name columns only (excludes first_name / last_name)
FULL_NAME_COL_RE = re.compile(
    r"(?<![A-Za-z])(?:employee[\s_]?name|emp[\s_]?name|full[\s_]?name|fullname"
    r"|(?<!first[\s_])(?<!last[\s_])name)(?![A-Za-z])",
    re.IGNORECASE
)

# Hint for DOB-like columns
DOB_COL_HINT_RE = re.compile(
    r"\b(dob|date\s*of\s*birth|birth\s*date|birthdate)\b",
    re.IGNORECASE
)

BANK_ACCT_COL_HINT_RE = re.compile(
    r"\b(bank\s*account|account\s*(?:no|number|num)|acct\s*(?:no|num))\b",
    re.IGNORECASE
)

# Hint for city columns
CITY_COL_HINT_RE = re.compile(
    r"\b(city|town|location|district)\b",
    re.IGNORECASE
)

# Normalize token for cross-type dedup
def _norm_token(pi_type: str, value: str) -> str:
    v = (value or "").strip()
    if pi_type.lower() in {"email"}:
        return v.lower()
    if pi_type.lower() in {"phone","ssn number","account number","credit/debit card","us passport number","company id","mrn","insurance id"}:
        return re.sub(r"\W+", "", v).upper()
    if pi_type.lower() == "date of birth":
        return v  # already ISO in our flow
    return re.sub(r"\s+", " ", v).strip()

TYPE_PRIORITY = {
    "Name, email": 1,
    "Name, Date of Birth": 1,
    "Name, Company ID": 1,
    "Name, Phone": 1,
    "Name, Address": 1,
    "Name, City": 1,
    "Name, State": 1,
    "Name, Gender": 1,
    "email": 2,
    "Date of Birth": 2,
    "SSN Number": 3,
    "Credit/Debit Card": 4,
    "US Passport Number": 5,
    "Driving License": 6,  # 'Driving License (ST)' maps to this base for priority comparison
    "Company ID": 7,
    "Account Number": 8,
    "Phone": 9,
    "Address": 10,
    "MRN": 11,
    "Insurance ID": 12,
}

def _base_type(pi_type: str) -> str:
    return "Driving License" if pi_type.startswith("Driving License") else pi_type

def find_ssn_in_text(text: str) -> list[str]:
    out = []
    # strict pattern: \b\d{3}-\d{2}-\d{4}\b
    out.extend(re.findall(r"\b\d{3}-\d{2}-\d{4}\b", text))
    # more flexible match when SSN keyword is nearby
    for m in SSN_NEARBY_FLEX_RE.finditer(text):
        out.append(f"{m.group(1)}-{m.group(2)}-{m.group(3)}")
    return list(dict.fromkeys(out))

def first_valid_dob_iso(text: str) -> str | None:
    raw_hits: list[str] = []
    for cp in DOB_PATTERNS:
        raw_hits.extend(m.group(0) for m in cp.finditer(text))
    for raw in raw_hits:
        dts = parse_multiple_dates(raw)
        for dt in dts:
            if is_probable_dob(dt):
                return dt.strftime("%Y-%m-%d")
    return None

def group_block_entities(block_text: str) -> dict[str, list[str]]:
    """
    Return entity buckets for names, emails, DOB, company IDs, and all other
    standalone PII types (SSN, passport, driving license, phone, address, etc.).
    """
    data: dict[str, list[str]] = {}
    t = block_text or ""
    if not t.strip():
        return data

    # Names
    # For pipe/tab-delimited table rows, try grabbing the first titlecase token cluster
    # if spaCy missed it (e.g. "Luna White | …")
    names = extract_person_names(t, max_names=5)

    if not names:
        _tbl_name_re = re.compile(
            r"^([A-Z][a-z]{1,15}(?:\s+[A-Z][a-z]{1,15}){1,2})\s*(?:\||,|\t|$)"
        )
        m = _tbl_name_re.match(t.strip())
        if m:
            cand = m.group(1).strip()
            toks = [x.upper() for x in cand.split()]
            if not any(x in PERSON_STOP_TOKENS for x in toks) and not _is_bad_label(cand):
                names = [cand]

    if names:
        data["name"] = names

    # Generic regex types (except SSN, handled below)
    for pi_type, pattern in pi_PATTERNS.items():
        if pi_type.lower() == "ssn number":
            continue
        for m in re.finditer(pattern, t):
            data.setdefault(pi_type, []).append(m.group(0))

    # US Passport (keyword-gated)
    if "US Passport Number" in data:
        data["US Passport Number"] = [p for p in data["US Passport Number"] if is_passport_token(p, t)]
        if not data["US Passport Number"]:
            data.pop("US Passport Number", None)

    # SSN (strict + flex)
    ssns = find_ssn_in_text(t)
    if ssns:
        data.setdefault("SSN Number", []).extend(ssns)

    # DOB — try all patterns including year-first variants
    dob_iso = first_valid_dob_iso(t)
    if dob_iso:
        data.setdefault("Date of Birth", []).append(dob_iso)
        for cp in DOB_PATTERNS:
            for m in cp.finditer(t):
                raw = m.group(0)
                data.setdefault("Date of Birth Raw", []).append(raw)

    # Driving License (context-gated)
    for state, full_value, core in iter_dl_matches(t, col_name=None, require_context=True):
        data.setdefault(f"Driving License ({state})", []).append(full_value)

    # MRN / Insurance
    for m in MRN_RE.finditer(t):
        data.setdefault("MRN", []).append(m.group(1) or m.group(2) or "")
    for m in INSURANCE_ID_RE.finditer(t):
        data.setdefault("Insurance ID", []).append(m.group(1) or m.group(2) or "")

    # Company ID (positive context + no negative context)
    for m in ALNUM_COMPANY_ID_RE.finditer(t):
        token = m.group(0)
        if is_company_id(token, context=t):
            data.setdefault("Company ID", []).append(token)
    for m in COMPANY_ID_RE_NUM.finditer(t):
        token = m.group(0)
        if is_company_id(token, context=t):
            data.setdefault("Company ID", []).append(token)

    # Account Number — only keep if banking keywords are present in context
    # Override: if "Account Number" was matched by regex but no context, drop it
    if "Account Number" in data:
        if not _has_account_context(t):
            data.pop("Account Number", None)
        else:
            # Remove Indian phone-shaped values unless explicitly labelled Account
            _ind_re = re.compile(r"^[6-9]\d{9}$")
            data["Account Number"] = [
                v for v in data["Account Number"]
                if not passes_luhn(v) and (
                    not _ind_re.match(re.sub(r"\D", "", v)) or _is_explicit_account(v, t)
                )
            ]
            if not data["Account Number"]:
                data.pop("Account Number", None)

    # Address — street pattern + suite-lead pattern
    flat = re.sub(r"[\r\n\t]+", " ", t)
    raw_addrs = []
    for m in re.findall(ADDRESS_PATTERN, flat, flags=re.IGNORECASE):
        v = m.strip()
        if v:
            raw_addrs.append(v)
    for m in ADDRESS_SUITE_PATTERN.finditer(flat):
        v = m.group(0).strip()
        if v and len(v) > 8:
            raw_addrs.append(v)
    # Drop any address that's a substring of a longer detected one
    raw_addrs_sorted = sorted(set(raw_addrs), key=len, reverse=True)
    deduped_addrs = []
    for addr in raw_addrs_sorted:
        if not any(addr != longer and addr in longer for longer in deduped_addrs):
            deduped_addrs.append(addr)
    if deduped_addrs:
        data["Address"] = deduped_addrs

    # Luhn-valid account numbers get re-tagged as credit cards
    for v in list(data.get("Account Number", [])):
        if passes_luhn(v):
            data.setdefault("Credit/Debit Card", []).append(v)
            data["Account Number"].remove(v)

    # Phone — filter to plausible format
    if "Phone" in data:
        data["Phone"] = [p for p in data["Phone"] if is_plausible_phone(p)]

    # Indian Phone — skip numbers already captured as account/routing (account wins when IFSC/bank context present)
    if "Indian Phone" in data:
        acct_digits = {re.sub(r"\D", "", v) for v in data.get("Account Number", [])}
        acct_digits |= {re.sub(r"\D", "", v) for v in data.get("Routing Number", [])}
        data["Indian Phone"] = [
            v for v in data["Indian Phone"]
            if re.sub(r"\D", "", v) not in acct_digits
        ]
        if not data["Indian Phone"]:
            data.pop("Indian Phone", None)

    # Gender
    gender_vals = detect_gender_values(t)
    if gender_vals:
        data["Gender"] = gender_vals

    # State — extract from text
    state_raw = detect_states(t, unique_only=True)
    if state_raw:
        data["State"] = state_raw if isinstance(state_raw, list) else [state_raw]

    # City — extract from addresses by looking for a city token before a state abbreviation/ZIP
    _city_from_addr_re = re.compile(
        r"\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,2}),\s*[A-Z]{2}(?:\s+\d{5,6})?\b"
    )
    _city_vals: List[str] = []
    for _addr in data.get("Address", []):
        for _m in _city_from_addr_re.finditer(_addr):
            _city_candidate = _m.group(1).strip()
            # Skip tokens that are known stop words (state names, common labels)
            # note: Indian city names like Bengaluru are intentionally excluded
            # because they are valid city values even if they appear in PERSON_STOP_TOKENS
            _city_stop_tokens = {
                "CALIFORNIA","TEXAS","FLORIDA","OHIO","GEORGIA","MICHIGAN","ILLINOIS",
                "PENNSYLVANIA","ARIZONA","INDIANA","TENNESSEE","MISSOURI","MARYLAND",
                "WISCONSIN","COLORADO","MINNESOTA","CAROLINA","VIRGINIA","NEVADA",
                "LOUISIANA","KENTUCKY","CONNECTICUT","OREGON","OKLAHOMA","IOWA",
                "MISSISSIPPI","ARKANSAS","UTAH","KANSAS","HAWAII","NEBRASKA","IDAHO",
                "MAINE","HAMPSHIRE","RHODE","ISLAND","MONTANA","DELAWARE","WYOMING",
                "ALASKA","VERMONT","DAKOTA","WASHINGTON",
                "FIELD","VALUE","SECTION","DOCUMENT","SYNTHETIC","NOTE","DISCLAIMER",
            }
            if _city_candidate.upper() not in _city_stop_tokens:
                _city_vals.append(_city_candidate)
    # also scan raw text for City, ST ZIP patterns
    for _m in _city_from_addr_re.finditer(t):
        _city_candidate = _m.group(1).strip()
        _city_stop_tokens = {
            "CALIFORNIA","TEXAS","FLORIDA","OHIO","GEORGIA","MICHIGAN","ILLINOIS",
            "PENNSYLVANIA","ARIZONA","INDIANA","TENNESSEE","MISSOURI","MARYLAND",
            "WISCONSIN","COLORADO","MINNESOTA","CAROLINA","VIRGINIA","NEVADA",
            "LOUISIANA","KENTUCKY","CONNECTICUT","OREGON","OKLAHOMA","IOWA",
            "MISSISSIPPI","ARKANSAS","UTAH","KANSAS","HAWAII","NEBRASKA","IDAHO",
            "MAINE","HAMPSHIRE","RHODE","ISLAND","MONTANA","DELAWARE","WYOMING",
            "ALASKA","VERMONT","DAKOTA","WASHINGTON",
            "FIELD","VALUE","SECTION","DOCUMENT","SYNTHETIC","NOTE","DISCLAIMER",
        }
        if _city_candidate.upper() not in _city_stop_tokens and _city_candidate not in _city_vals:
            _city_vals.append(_city_candidate)
    if _city_vals:
        data["City"] = list(dict.fromkeys(_city_vals))  # deduped, order preserved

    # Deduplicate per bucket
    for k, vals in list(data.items()):
        seen = set(); uniq = []
        for x in vals:
            if x not in seen:
                seen.add(x); uniq.append(x)
        data[k] = uniq

    return data

# Dedup helpers
def _base_type(pi_type: str) -> str:
    """Normalize PI Type for de-dup purposes (e.g., all 'Driving License (ST)' → 'Driving License')."""
    return "Driving License" if pi_type.startswith("Driving License") else pi_type

def _norm_value_for_key(base_type: str, value: str) -> str:
    """Normalize detected values for dedup keys: emails lowercased, numbers stripped of punctuation, dates as-is, everything else collapsed."""
    v = (value or "").strip()
    bt = base_type.lower()
    if bt == "email":
        return v.lower()
    if bt in {"phone", "ssn number", "account number", "credit/debit card",
              "us passport number", "company id", "mrn", "insurance id"}:
        return re.sub(r"\W+", "", v).upper()
    if bt == "date of birth":
        return v  # already normalized to ISO in your flow
    # default: collapse whitespace
    return re.sub(r"\s+", " ", v).strip()

def _add_record(records: List[Dict[str, Any]],
                seen: set[tuple],
                *,
                file_name: str,
                page_sheet: str,
                cell: str,
                pi_type: str,
                value: str,
                confidence: int) -> None:
    base = _base_type(pi_type)
    norm_val = _norm_value_for_key(base, value)
    # Pair types: dedup globally so the same pair isn't repeated from multiple window positions
    if pi_type.startswith("Name,"):
        key = (file_name, base, norm_val)
    else:
        # Standalone PII: include page so the same value on different pages both shows up
        base_page = re.sub(r"\s*\(Table\s+\d+\)\s*$", "", page_sheet, flags=re.IGNORECASE).strip()
        key = (file_name, base_page, base, norm_val)
    if key in seen:
        return
    seen.add(key)
    records.append({
        "File": file_name,
        "Page/Sheet": page_sheet,
        "Cell": cell,
        "PI Type": pi_type,
        "Detected Value": value,
        "Confidence (%)": confidence,
    })

def _find_best_name_for_target(names: list[str], block_text: str, target_value: str) -> str | None:
    """Pick the name closest to the target in block_text, preferring names that appear before it."""
    if not names or not target_value:
        return names[0] if names else None

    t = block_text or ""

    try:
        tgt_idx = t.index(target_value)
    except ValueError:
        tgt_idx = -1

    scored: list[tuple[int, int, str]] = []  # (penalty, distance, name)
    for n in names:
        try:
            n_idx = t.index(n)
        except ValueError:
            # if not found verbatim (OCR quirks), fall back to distance from start
            n_idx = 0

        # Names before the target get penalty 0, names after get penalty 1
        penalty = 0 if (tgt_idx >= 0 and n_idx <= tgt_idx) else 1
        dist = abs((tgt_idx if tgt_idx >= 0 else 0) - n_idx)
        scored.append((penalty, dist, n))

    # sort by penalty then distance
    scored.sort(key=lambda x: (x[0], x[1]))
    return scored[0][2] if scored else (names[0] if names else None)

def emit_grouped_records(file_name: str,
                         page_or_sheet: str,
                         cell: str,
                         block_text: str,
                         pair_only: bool = False) -> list[dict]:
    """Emit name-paired records (Name+email, Name+DOB, Name+CompanyID etc.) using proximity within the block."""
    results: list[dict] = []
    bucket = group_block_entities(block_text)
    names = bucket.pop("name", [])
    if not names:
        return [] if pair_only else []

    for email in bucket.get("email", []):
            best = _find_best_name_for_target(names, block_text, email)
            if best:
                results.append({
                "File": file_name,
                "Page/Sheet": page_or_sheet,
                "Cell": cell,
                "PI Type": "Name, email",
                "Detected Value": f"{best}, {email}",
                "Confidence (%)": 95,
            })

    # Pair DOB
    for dob in bucket.get("Date of Birth", []):
        best = _find_best_name_for_target(names, block_text, dob)
        if best:
            results.append({
                "File": file_name,
                "Page/Sheet": page_or_sheet,
                "Cell": cell,
                "PI Type": "Name, Date of Birth",
                "Detected Value": f"{best}, {dob}",
                "Confidence (%)": 95,
            })

    # Pair Company ID
    for cid in bucket.get("Company ID", []):
        best = _find_best_name_for_target(names, block_text, cid)
        if best:
            results.append({
                "File": file_name,
                "Page/Sheet": page_or_sheet,
                "Cell": cell,
                "PI Type": "Name, Company ID",
                "Detected Value": f"{best}, {cid}",
                "Confidence (%)": 95,
            })

    for phone in bucket.get("Phone", []):
        best = _find_best_name_for_target(names, block_text, phone)
        if best:
            results.append({
                "File": file_name,
                "Page/Sheet": page_or_sheet,
                "Cell": cell,
                "PI Type": "Name, Phone",
                "Detected Value": f"{best}, {phone}",
                "Confidence (%)": 90,
            })

    for phone in bucket.get("Indian Phone", []):
        best = _find_best_name_for_target(names, block_text, phone)
        if best:
            results.append({
                "File": file_name,
                "Page/Sheet": page_or_sheet,
                "Cell": cell,
                "PI Type": "Name, Phone",
                "Detected Value": f"{best}, {phone}",
                "Confidence (%)": 90,
            })

    for addr in bucket.get("Address", []):
        best = _find_best_name_for_target(names, block_text, addr)
        if best:
            results.append({
                "File": file_name,
                "Page/Sheet": page_or_sheet,
                "Cell": cell,
                "PI Type": "Name, Address",
                "Detected Value": f"{best}, {addr}",
                "Confidence (%)": 88,
            })

    # City pairs
    for city in bucket.get("City", []):
        best = _find_best_name_for_target(names, block_text, city)
        if best:
            results.append({
                "File": file_name,
                "Page/Sheet": page_or_sheet,
                "Cell": cell,
                "PI Type": "Name, City",
                "Detected Value": f"{best}, {city}",
                "Confidence (%)": 85,
            })

    # Normalize states — convert full names to postal codes to avoid NY + New York duplicates
    _state_list = bucket.get("State", [])
    _seen_state_codes: set = set()
    _normalized_states: list = []
    for s in _state_list:
        # Convert full state name to postal code if possible
        _code = US_STATES.get(s) or US_STATES.get(s.title()) or s
        # Also check if s is already a 2-letter code in the values dict
        if len(s) > 2:  # full name like "New York"
            _code = US_STATES.get(s, s)
        else:
            _code = s  # already a postal code
        if _code not in _seen_state_codes:
            _seen_state_codes.add(_code)
            _normalized_states.append(_code)

    for state in _normalized_states:
        best = _find_best_name_for_target(names, block_text, state)
        if best:
            results.append({
                "File": file_name,
                "Page/Sheet": page_or_sheet,
                "Cell": cell,
                "PI Type": "Name, State",
                "Detected Value": f"{best}, {state}",
                "Confidence (%)": 80,
            })

    for gender in bucket.get("Gender", []):
        best = _find_best_name_for_target(names, block_text, gender)
        if best:
            results.append({
                "File": file_name,
                "Page/Sheet": page_or_sheet,
                "Cell": cell,
                "PI Type": "Name, Gender",
                "Detected Value": f"{best}, {gender}",
                "Confidence (%)": 85,
            })

    return results

def passes_luhn(s: str) -> bool:
    digits = re.sub(r"\D", "", s)
    if len(digits) < 12:  # cards are typically 13-19 digits; 12 is a safe OCR floor
        return False
    total = 0
    rev = digits[::-1]
    for i, ch in enumerate(rev):
        d = ord(ch) - 48
        if i % 2 == 1:
            d = d * 2
            if d > 9:
                d -= 9
        total += d
    return total % 10 == 0

def is_plausible_phone(s: str) -> bool:
    digits = re.sub(r"\D", "", s)
    # Require at least 10 digits and at least one phone-like symbol/space
    return len(digits) >= 10 and re.search(r"[()\-\+\s]", s) is not None

DL_CONTEXT_RE = re.compile(
    r"\b(driver'?s?|driving)\b|\blic(?:\.|ense)?\b|\bdl\b|\bd[.\s]?/?[\s]*l\.?\b"
    r"|driver(?:s)?licen[cs]e"   # compound: DriverLicense, DriversLicense
    r"|driving[\s_]?licen[cs]e", # compound: DrivingLicense
    re.IGNORECASE
)

def has_dl_context(text: str, col_name: str | None = None) -> bool:
    """True if DL-like keywords appear in the text or column header."""
    def _hit(s: str | None) -> bool:
        return bool(s and DL_CONTEXT_RE.search(s))
    return _hit(text) or _hit(str(col_name) if col_name is not None else None)

# Safe string conversion for cell values
def _to_text(val: Any) -> str:
    try:
        if pd.isna(val):
            return ""
    except Exception:
        pass
    return str(val).strip() if val is not None else ""

def _dl_core_len_ok(state: str, core: str) -> bool:
    """Validate DL core length using min/max tables."""
    normalized_core = re.sub(r"[^A-Za-z0-9]", "", core or "")
    normalized_core = re.sub(rf"^{state}", "", normalized_core, flags=re.IGNORECASE)
    min_len = US_DL_MIN_LENGTH.get(state, 1)
    max_len = US_DL_MAX_LENGTH.get(state, len(normalized_core))
    return min_len <= len(normalized_core) <= max_len

def iter_dl_matches(text: str, col_name: str | None = None, require_context: bool = True):
    """Yield (state, full_value, core) for DL candidates. Longest match wins for overlapping spans. Values adjacent to a passport keyword are skipped."""
    if not text:
        return
    if require_context and not has_dl_context(text, col_name):
        return

    # Pre-collect passport keyword spans to suppress DL matches that sit next to them
    _PASSPORT_KW_RE = re.compile(r"\bpassport\b", re.IGNORECASE)
    passport_spans = [(m.start(), m.end()) for m in _PASSPORT_KW_RE.finditer(text)]

    # Check if col_name is explicitly a DL/license column
    _col_str = str(col_name).lower() if col_name else ""
    _is_dl_col = bool(re.search(r'\b(?:driver|driving|license|licence|dl)\b', _col_str))

    def _near_passport(start: int, end: int, value: str) -> bool:
        """True if the value sits immediately after a passport keyword (within 15 chars)."""
        for ps, pe in passport_spans:
            # Only suppress if the match starts right after the passport keyword
            if 0 <= (start - pe) <= 15:
                return True
        return False

    # Collect all raw matches, then deduplicate by span
    all_matches = []  # (start, end, state, full_value, core)
    for state, pattern in US_DL_PATTERNS.items():
        # Pure-digit DL states are too ambiguous — only emit when
        # an explicit DL/driver/license column header is present
        if state in _PURE_DIGIT_DL_STATES and not _is_dl_col:
            continue
        for m in re.finditer(pattern, text, flags=re.IGNORECASE):
            full_value = m.group(0)
            core = (m.groupdict().get("core") if hasattr(m, "groupdict") else None) or full_value
            if not _dl_core_len_ok(state, core):
                continue
            # skip if this value sits immediately after a passport keyword
            if _near_passport(m.start(), m.end(), full_value):
                continue
            all_matches.append((m.start(), m.end(), state, full_value, core))

    # Sort by length descending — longest/most-specific match wins
    all_matches.sort(key=lambda x: -(x[1] - x[0]))

    emitted_spans: list[tuple[int, int]] = []
    for start, end, state, full_value, core in all_matches:
        # Skip if this span overlaps an already-emitted (longer) match
        if any(not (end <= es or start >= ee) for es, ee in emitted_spans):
            continue
        emitted_spans.append((start, end))
        yield state, full_value, core

# Fallback person name detector (used when spaCy isn't available) and blacklist filters
PERSON_HEUR_RE = re.compile(r"\b([A-Z][a-z]+(?:[ \t]+[A-Z][a-z]+){1,2})(?=[ \t]*(?:[,;:\.\-\u2013\u2014(]|\s|$))")

# Tokens printed on licenses/IDs that should never be treated as names
STOP_TOKENS = {
    # Document/label tokens
    "DRIVER", "DRIVERS", "DRIVING", "LICENSE", "ID", "DL", "NUMBER", "NO", "NUM",
    "ADDRESS", "ADDR", "STATE", "COUNTRY", "USA", "UNITED", "AMERICA",
    "CLASS", "SEX", "ENDORSEMENTS", "ENDORSEMENT", "ENDORSE", "ENDO", "RSTR", "RESTRICTIONS", "RESTRICTION",
    "DONOR", "ORGAN", "EYE", "HAIR", "HEIGHT", "HGT", "WEIGHT", "WGT", "SIGN", "SIGNATURE",
    # Dates/metadata
    "DOB", "BIRTH", "DATE", "ISS", "ISSUE", "ISSUED", "EXP", "EXPIRES", "EXPIRY",
    # OCR noise and test data words
    "TEST", "DATA", "TEST DATA", "SAMPLE", "SPECIMEN", "BARCODE",
    # Section/document header words sometimes mistagged as PERSON
    "SCENARIO", "ONBOARDING", "PAYMENT", "DISPUTE", "MEDICAL", "PRE", "AUTHORIZATION",
    "WIRE", "TRANSFER", "REQUEST", "REQUESTER", "CORPORATE", "SYNTHETIC", "EMBEDDED",
    "DEAR", "WRITING", "CONFIRM", "START", "SESSION", "TRANSACTION", "UNAUTHORIZED",
    "EMERGENCY", "VERIFICATION", "ORIGINATED",
    # Direction/role prefixes that sometimes get merged with names
    "FROM", "TO", "SENDER", "RECEIVER", "RECIPIENT", "SUBJECT", "RE",
    # Section titles
    "FRAUD", "INVESTIGATION", "LOAN", "APPLICATION", "ADDITIONAL", "RECORDS",
    "BACKGROUND", "CHECK", "EMPLOYEE", "HEALTHCARE", "FINANCE", "GLOBAL",
    # Column/category labels that can trail names in Excel cells
    "IDENTITY", "CATEGORY", "PROOF", "TYPE", "SOURCE", "STATUS", "LABEL",
    "FIELD", "VALUE", "RECORD", "ENTRY", "ITEM", "FORMAT", "DETAIL",
    "DESCRIPTION", "REMARK", "COMMENT", "VERSION", "REVISION",
    "TABLE", "PAGE", "SECTION", "BLOCK", "FORM", "TEMPLATE", "REPORT",
    "SUMMARY", "OVERVIEW", "HEADER", "COLUMN", "ROW",
    "SCANNED", "DOCUMENT", "IMAGE", "STRUCTURED", "SELECTABLE",
    "TYPED", "DUMMY", "COMPARISON", "DRAFT", "REVIEW", "PREVIEW",
    # Words that trail names in free text (| Gender: Male | Company ID:)
    "GENDER", "FEMALE", "MALE", "COMPANY", "EMPLOYEE", "STAFF", "BADGE", "PAYROLL",
    "PII", "FULL", "SET", "TEST", "COMPREHENSIVE", "COMPLETE",
    # Medical/healthcare words often misdetected as names
    "ADMISSION", "DISCHARGE", "PATIENT", "DIAGNOSIS", "TREATMENT", "PROCEDURE",
    "REFERRAL", "PRESCRIPTION", "ATTENDING", "RESIDENT", "PHYSICIAN", "NURSING",
    "INSURANCE", "CLAIM", "BILLING", "AUTHORIZATION", "PREAUTH",
    # Finance words
    "WITHDRAWAL", "DEPOSIT", "SETTLEMENT", "CLEARANCE", "PRINCIPAL",
    "INTEREST", "BALANCE", "LEDGER", "BENEFICIARY",
}

# Common verbatim phrases that appear on IDs
STOP_PHRASES = {
    "DRIVER LICENSE", "DRIVER'S LICENSE", "DATE OF BIRTH", "TEST DATA", "DRIVING LICENSE",
    "WIRE TRANSFER", "WIRE TRANSFER REQUEST", "MEDICAL PRE", "MEDICAL PRE-AUTHORIZATION",
    "PAYMENT DISPUTE", "ONBOARDING EMAIL",
    "FRAUD INVESTIGATION", "LOAN APPLICATION", "ADDITIONAL RECORDS", "BACKGROUND CHECK",
    "EMPLOYEE BACKGROUND", "EMPLOYEE BACKGROUND CHECK",
}

# Name extraction and cleaning

# Tokens, headers, and cities that should never be treated as names
PERSON_STOP_TOKENS = {
    "HELLO","HI","INTAKE","EMAIL","SUPPORT","TICKET","PAYMENT","BANK","ACCOUNT","UPDATE","VOTER","APPLICANT","POLICYHOLDER","CLAIMANT","INSURED","NOMINEE","CUSTOMER","CLIENT","AADHAAR","AADHAR","PAN","IFSC","UID","UIN",
    "ROUTING","SCENARIO","TRAVEL","ITINERARY","NOTE","DISCLAIMER","CONTACT","TABLE","EMBEDDED",
    # Section / document header words that appear as labels in test documents
    "FIELD","VALUE","SECTION","SELECTABLE","TYPED","STRUCTURED","SCANNED","DOCUMENT","IMAGE","FREE","FORM","BLOCK","COMPARISON","OCR","SPREADSHEET","DUMMY","SYNTHETIC",
    # US city tokens (multi-word cities handled by _is_us_city_fragment)
    "MIAMI","SEATTLE","RALEIGH","ATLANTA","SAN","FRANCISCO","DALLAS","DENVER","DETROIT",
    "ANGELES","CHICAGO","HOUSTON","PHOENIX","PHILADELPHIA","ANTONIO","DIEGO","JOSE",
    "AUSTIN","JACKSONVILLE","COLUMBUS","CHARLOTTE","INDIANAPOLIS","PORTLAND","MEMPHIS",
    "LOUISVILLE","BALTIMORE","BOSTON","NASHVILLE","OKLAHOMA","ALBUQUERQUE","TUCSON",
    "FRESNO","SACRAMENTO","MESA","KANSAS","OMAHA","CLEVELAND","MINNEAPOLIS","WICHITA",
    "ARLINGTON","COLORADO","SPRINGS","TAMPA","ORLEANS","HONOLULU","ANAHEIM","AURORA",
    "CORPUS","CHRISTI","RIVERSIDE","LEXINGTON","HENDERSON","STOCKTON","ANCHORAGE","NEWARK",
    "PITTSBURGH","GREENSBORO","LINCOLN","CINCINNATI","TOLEDO","JERSEY","ORLANDO",
    "PLANO","CHULA","VISTA","IRVINE","LAREDO","MADISON","DURHAM","LUBBOCK","GARLAND",
    "WINSTON","SALEM","NORFOLK","SCOTTSDALE","CHANDLER","BATON","ROUGE","HIALEAH",
    "FREMONT","RICHMOND","BAKERSFIELD","GILBERT","BIRMINGHAM","ROCHESTER","SPOKANE",
    "DES","MOINES","MONTGOMERY","GLENDALE","TACOMA","AKRON","SHREVEPORT","HUNTINGTON",
    "LITTLE","ROCK","COLUMBIA","AUGUSTA","GRAND","RAPIDS","SALT","LAKE","CITY",
    "TALLAHASSEE","HUNTSVILLE","WORCESTER","KNOXVILLE","PROVIDENCE","BROWNSVILLE",
    "SANTA","ANA","BUFFALO","FORT","WORTH","EL","PASO","NEW","YORK","LOS","WASHINGTON",
    # US state names (can be misdetected as person names)
    "CALIFORNIA","TEXAS","FLORIDA","OHIO","GEORGIA","MICHIGAN","ILLINOIS","PENNSYLVANIA",
    "ARIZONA","INDIANA","TENNESSEE","MISSOURI","MARYLAND","WISCONSIN","COLORADO",
    "MINNESOTA","CAROLINA","VIRGINIA","NEVADA","LOUISIANA","KENTUCKY","CONNECTICUT",
    "OREGON","OKLAHOMA","IOWA","MISSISSIPPI","ARKANSAS","UTAH","KANSAS","NEVADA",
    "HAWAII","NEBRASKA","IDAHO","MAINE","HAMPSHIRE","RHODE","ISLAND","MONTANA",
    "DELAWARE","WYOMING","ALASKA","VERMONT","DAKOTA",
    # Indian location words often mistagged as PERSON
    "NAGAR","SALAI","COLONY","HILLS","LINES","LAKE","LAYOUT","EXTENSION","VIHAR","ENCLAVE",
    "SECTOR","BLOCK","PHASE","CROSS","MAIN","CHOWK","MARG",
    # Indian cities commonly misdetected as names
    "MUMBAI","DELHI","BENGALURU","BANGALORE","CHENNAI","HYDERABAD","PUNE","KOLKATA",
    "AHMEDABAD","JAIPUR","LUCKNOW","SURAT","BHOPAL","NAGPUR","PATNA","INDORE",
    # Location prefixes
    "CIVIL","JUBILEE","LAJPAT","NEHRU","ANNA","BRIGADE","MARINE","PARK","MG",
    # Email domain words sometimes extracted as names
    "EXAMPLE","GMAIL","YAHOO","HOTMAIL","OUTLOOK",
}

# Multi-word US city fragments that spaCy sometimes tags as PERSON
_US_CITY_FRAGMENTS = {
    "LOS ANGELES", "NEW YORK", "SAN FRANCISCO", "SAN DIEGO", "SAN JOSE", "SAN ANTONIO",
    "NEW ORLEANS", "LAS VEGAS", "EL PASO", "FORT WORTH", "COLORADO SPRINGS",
    "CORPUS CHRISTI", "SALT LAKE CITY", "BATON ROUGE", "CHULA VISTA", "GRAND RAPIDS",
    "LITTLE ROCK", "WINSTON SALEM", "SANTA ANA", "FORT LAUDERDALE", "NORTH LAS VEGAS",
    "JERSEY CITY", "OKLAHOMA CITY", "KANSAS CITY",
}

def _is_us_city_fragment(name: str) -> bool:
    """Return True if the candidate name matches a known US city (multi-word)."""
    return name.upper() in _US_CITY_FRAGMENTS

# Contextual name patterns — lookahead stops at punctuation/parens/period
_CTX_NAME_PATTERNS = [
    re.compile(r"\bmy\s+name\s+is\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})(?=\s*(?:[,;:\.\-\u2013\u2014(]|$))", re.IGNORECASE),
    re.compile(r"\bpatient(?:\s*:)?\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})(?=\s*(?:[,;:\.\-\u2013\u2014(]|$))",   re.IGNORECASE),
    re.compile(r"\btraveler(?:\s*:)?\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})(?=\s*(?:[,;:\.\-\u2013\u2014(]|$))",  re.IGNORECASE),
    re.compile(r"\b(?:applicant|voter|policy\s*holder|employee)(?:[:\s]+)([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})(?=\s*(?:[,;\.\s]|$))", re.IGNORECASE),
    re.compile(r"\b(?:customer|requester|requestor)(?:\s*:)?\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})(?=\s*(?:[,;\.\-]|$))", re.IGNORECASE),
    # Direction/role prefixes — stop at comma, pipe, or period
    re.compile(r"\b(?:from|sender)(?:\s*:)?\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})(?=\s*(?:[,;\.\|\-]|$))", re.IGNORECASE),
    re.compile(r"\b(?:to|receiver|recipient)(?:\s*:)?\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})(?=\s*(?:[,;\.\|\-]|$))", re.IGNORECASE),
    re.compile(r"\bsubject(?:\s*:)?\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})(?=\s*(?:[,;\.\|\-]|$))", re.IGNORECASE),
    re.compile(r"\bName(?:[:\s]+)([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})(?=\s*(?:[,;|\.\s]|$))", re.IGNORECASE),
]


# Strips label prefixes at the start (e.g. 'Patient:', 'Name:')
NAME_PREFIX_RE  = re.compile(
    r"^(?:patient|traveler|name|voter|applicant|policy\s*holder|requester|requestor|"
    r"employee|contact|claimant|insured|beneficiary|nominee|customer|client|"
    r"from|to|sender|receiver|recipient|subject|re)(?:\s*:)?\s+"
    r"|^(?:my\s+name\s+is|is|was|are)\s+",
    re.IGNORECASE
)

# Remove trailing parenthetical chunks
NAME_PAREN_TRAILER_RE = re.compile(r"\s*\([^)]*\)\s*$")

# Remove trailers starting with MRN/DOB/SSN/DL/Passport etc.
NAME_BAD_TRAILERS_RE = re.compile(
    r"(?:[,;:\-\u2013\u2014]\s*)?(?:"
    r"mrn|medical\s*record|dob|d\.o\.b\.|date\s*of\s*birth|ssn|social\s*security|dl|driver'?s?\s*license|passport"
    r")\b.*$",
    re.IGNORECASE
)

# Drop any trailing label token from a name
NAME_BAD_TAIL_TOKEN_RE = re.compile(
    r"\b(?:mrn|dob|ssn|dl|passport|policy|account|routing|"
    r"aadhaar|aadhar|pan|ifsc|voter|id|uid|uin|gstin|cin)\.?$",
    re.IGNORECASE
)

# Reject the name if any of these tokens appear
NAME_REJECT_TOKEN_RE = re.compile(
    r"\b(?:mrn|medical\s*record|dob|d\.o\.b\.|date\s*of\s*birth|ssn|social\s*security|"
    r"dl|driver\'?s?\s*license|passport|aadhaar|aadhar|pan\s*card|ifsc|voter\s*id|"
    r"account\s*number|bank\s*account|insurance\s*id|policy\s*no)\b",
    re.IGNORECASE
)

def _ctx_names(text: str) -> list[str]:
    out = []
    for rx in _CTX_NAME_PATTERNS:
        for m in rx.finditer(text or ""):
            out.append(m.group(1).strip())
    seen = set(); uniq = []
    for n in out:
        if n not in seen:
            seen.add(n); uniq.append(n)
    return uniq

def _clean_name(n: str) -> str:
    if not n:
        return ""
    # 1) strip prefixes
    n = NAME_PREFIX_RE.sub("", n).strip()
    # Remove trailing parenthetical chunks
    # run twice to be safe with nested OCR artifacts
    for _ in range(2):
        n = NAME_PAREN_TRAILER_RE.sub("", n).strip()
    # 3) remove MRN/DOB/etc. trailers
    n = NAME_BAD_TRAILERS_RE.sub("", n).strip()
    # Drop trailing label token if any
    n = NAME_BAD_TAIL_TOKEN_RE.sub("", n).strip()
    # 5) collapse whitespace
    n = re.sub(r"\s+", " ", n).strip()
    return n

def extract_person_names(text: str, max_names: int = 20) -> list[str]:
    """Prefer contextual names, then spaCy PERSON entities, then heuristic. Cleans label prefixes/trailers and rejects address fragments, cities, and anything with digits."""
    if not text or not text.strip():
        return []
    normalized = re.sub(r"[^A-Za-z0-9' \-\n]", " ", text)

    candidates: list[str] = []

    # contextual — run on both raw and normalized text
    candidates.extend(_ctx_names(text))
    candidates.extend(_ctx_names(normalized))

    # spaCy PERSON entities
    nlp = get_nlp()
    if nlp is not None:
        try:
            doc = nlp(normalized)
            candidates.extend([ent.text.strip() for ent in doc.ents if ent.label_ == "PERSON"])
        except Exception:
            pass

    # heuristic fallback
    candidates.extend(detect_person_names_simple(normalized))

    # Clean + filter + de-dup
    seen = set()
    cleaned: list[str] = []
    for raw in candidates:
        n = _clean_name(raw)
        if not n or len(n) < 3:
            continue
        if re.search(r"\d", n):               # any digit means it's not a name
            continue
        # Only accept uppercase-starting single words
        if len(n.split()) == 1 and not n[0].isupper():
            continue
        # in free text, single words are not reliable names without context
        # Single words are not reliable names without context
        if _is_bad_label(n):                  # your label filter (e.g., SEX/CLASS/TEST)
            continue

        # Avoid headers/cities/etc. (you already maintain PERSON_STOP_TOKENS)
        toks = [t.upper() for t in re.findall(r"[A-Za-z]+", n)]
        if any(t in PERSON_STOP_TOKENS for t in toks):
            continue

        # Reject multi-word US city names
        if _is_us_city_fragment(n):
            continue

        # reject address-like fragments or things in an address zone
        if _looks_like_address_fragment(n, normalized):
            continue

        if n not in seen:
            seen.add(n)
            cleaned.append(n)

    return cleaned[:max_names]

# Generic words we also don’t want as names (used by the simpler heuristic)
PERSON_BLACKLIST = {
    "January","February","March","April","May","June","July","August","September","October","November","December",
    "Employee","Employees","Company","Address","State","Page","Table","Row","Column","Sheet","Name","Names",
    "All","Insurance","Registration","Scenario","Policy","Voter","Applicant","Disclaimer","Note","Card","Phone","Aadhaar","Pan","Ifsc","Email","Mobile","Address","Bank","Account",
    "Synthetic","Fabricated","Testing","Verification","Update","Review","Contact","Embedded",
    "Following","Flagged","Reserved","Purposes","Functional","Non","Please","Hello",
    "Field","Value","Section","Selectable","Typed","Structured","Scanned","Document","Image","Free","Form","Block","Comparison","Ocr","Spreadsheet","Dummy",
    # Credit card brands
    "Visa","Mastercard","Amex","Discover","Maestro","Rupay","Diners",
    # Gender words
    "Male","Female","Gender","Transgender","Binary",
    # Document title words
    "Pii","Full","Set","Comprehensive","Complete","Test","Sample",
    # Other common FPs from Excel column values
    "Fake","Data","Only","Record","Entry","Info","Information",
    "Medical","Financial","Corporate","Global","Healthcare","India","Indian",
}

def _is_bad_label(name: str) -> bool:
    """
    Return True if the candidate looks like a label/field/value that should not be treated as a person name.
    """
    if not name:
        return True

    # Normalize to uppercase tokens (letters and apostrophes preserved)
    t = re.sub(r"[^A-Za-z'\s-]", " ", name).upper()
    tokens = [tok for tok in re.split(r"\s+", t) if tok]
    if not tokens:
        return True

    # Phrase ban
    phrase = " ".join(tokens)
    if phrase in STOP_PHRASES:
        return True

    # If 50% or more tokens are stop words, treat as a label
    bad_ratio = sum(1 for tok in tokens if tok in STOP_TOKENS) / len(tokens)
    if bad_ratio >= 0.5:
        return True

    # For 2-word candidates: reject if EITHER word is a stop token
    
    if len(tokens) == 2 and any(tok in STOP_TOKENS for tok in tokens):
        return True

    # For 3-word candidates: reject if the LAST word is a stop token
    
    # The first two words are a valid name but the trailing label word pollutes it
    if len(tokens) == 3 and tokens[-1] in STOP_TOKENS:
        return True

    # Reject single-word label tokens
    if len(tokens) == 1 and tokens[0] in STOP_TOKENS:
        return True

    # Column-header style: short/all-caps 12 tokens with no vowels (often OCR noise)
    if len(tokens) <= 2 and all(tok.isupper() for tok in tokens) and all(len(tok) <= 6 for tok in tokens):
        if not any(re.search(r"[AEIOU]", tok) for tok in tokens):
            return True

    return False

def detect_person_names_simple(text: str) -> list[str]:
    """
    Heuristic fallback: 2–3 capitalized words. Light filtering only.
    Applies _clean_name() before PERSON_BLACKLIST check so trailing label tokens
    (e.g. 'Aadhaar', 'PAN') are stripped before the candidate is evaluated.
    When a 3-word candidate's last word is a stop token (e.g. 'Priya Sharma Identity'),
    the last word is trimmed and the 2-word name ('Priya Sharma') is emitted instead.
    Final filtering is handled in extract_person_names().
    """
    if not text:
        return []
    out = []
    for m in PERSON_HEUR_RE.finditer(text):
        cand = m.group(1).strip()
        # Clean trailing label tokens BEFORE blacklist check
        cand = _clean_name(cand)
        if not cand:
            continue
        parts = cand.split()
        if any(p in PERSON_BLACKLIST for p in parts):
            continue
        if re.search(r"\d", cand):
            continue
        # If 3-word candidate ends with a stop token, trim it to first 2 words
        if len(parts) == 3 and parts[-1].upper() in STOP_TOKENS:
            cand = " ".join(parts[:2])
        out.append(cand)
    seen: set = set()
    unique = []
    for x in out:
        if x not in seen:
            seen.add(x)
            unique.append(x)
    return unique

# Address fragment filters
ADDRESS_TOKENS = {
    "ROAD","RD","STREET","ST","AVENUE","AVE","BOULEVARD","BLVD","LANE","LN","DRIVE","DR",
    "COURT","CT","CIRCLE","CIR","HIGHWAY","HWY","TERRACE","TER","PLACE","PL","SQUARE","SQ",
    "PARKWAY","PKWY","WAY","SUITE","STE","APARTMENT","APT","UNIT","FL","FLOOR","BLDG","BUILDING",
    "PO","BOX",
    # Indian address tokens  prevents location strings being treated as names
    "NAGAR","SALAI","COLONY","HILLS","LINES","LAKE","LAYOUT","VIHAR","ENCLAVE",
    "SECTOR","CROSS","CHOWK","MARG","PHASE","BLOCK",
}

ADDRESS_LEADER_RE = re.compile(r"\b(address\s*:?)\b", re.IGNORECASE)

def _looks_like_address_fragment(name: str, ctx: str) -> bool:
    """Reject if candidate contains street/unit words or sits in an 'Address' zone."""
    n = (name or "").strip()
    if not n:
        return False
    toks = [t.upper() for t in re.findall(r"[A-Za-z]+", n)]
    # reject if address tokens present
    if any(t in ADDRESS_TOKENS for t in toks):
        return True
    # reject if inside an address zone
    idx = (ctx or "").find(n)
    if idx >= 0:
        zone = (ctx[max(0, idx - 40): idx] or "")
        if ADDRESS_LEADER_RE.search(zone):
            return True
    return False

# US States
US_STATES = {
    "Alabama": "AL", "Alaska": "AK", "Arizona": "AZ", "Arkansas": "AR",
    "California": "CA", "Colorado": "CO", "Connecticut": "CT", "Delaware": "DE",
    "Florida": "FL", "Georgia": "GA", "Hawaii": "HI", "Idaho": "ID",
    "Illinois": "IL", "Indiana": "IN", "Iowa": "IA", "Kansas": "KS",
    "Kentucky": "KY", "Louisiana": "LA", "Maine": "ME", "Maryland": "MD",
    "Massachusetts": "MA", "Michigan": "MI", "Minnesota": "MN", "Mississippi": "MS",
    "Missouri": "MO", "Montana": "MT", "Nebraska": "NE", "Nevada": "NV",
    "New Hampshire": "NH", "New Jersey": "NJ", "New Mexico": "NM", "New York": "NY",
    "North Carolina": "NC", "North Dakota": "ND", "Ohio": "OH", "Oklahoma": "OK",
    "Oregon": "OR", "Pennsylvania": "PA", "Rhode Island": "RI", "South Carolina": "SC",
    "South Dakota": "SD", "Tennessee": "TN", "Texas": "TX", "Utah": "UT",
    "Vermont": "VT", "Virginia": "VA", "Washington": "WA", "West Virginia": "WV",
    "Wisconsin": "WI", "Wyoming": "WY", "District of Columbia": "DC"
}
STATE_NAME_RE = re.compile(r"\b(" + "|".join(re.escape(n) for n in US_STATES.keys()) + r")\b", re.IGNORECASE)
STATE_POSTAL_RE = re.compile(r"(?<![A-Za-z])(?<!\w)(" + "|".join(US_STATES.values()) + r")(?![A-Za-z])")

# skip high-ambiguity postal codes unless in address context
_AMBIGUOUS_STATE_CODES = {"ID", "IN", "OR", "OK", "ME", "OH", "AR", "CO", "DE", "HI", "IA", "KS", "KY", "LA", "MA", "MD", "MI", "MN", "MO", "MS", "MT", "NC", "ND", "NE", "NH", "NJ", "NM", "NV", "NY", "OR", "PA", "RI", "SC", "SD", "UT", "VA", "VT", "WA", "WI", "WV", "WY"}
_HIGH_AMBIGUITY_CODES = {"ID"}

def detect_states(text: str, unique_only: bool = False) -> List[str] | str:
    """Detect US states (returns comma-joined string by default).
    Two-letter postal codes that are highly ambiguous (e.g. 'ID' for Idaho)
    are only accepted when the surrounding context strongly suggests a state
    (e.g. preceded by a city name and comma, or inside an address pattern).
    """
    if not text:
        return "" if not unique_only else []
    results: List[str] = []
    for m in STATE_NAME_RE.finditer(text):
        results.append(next(name for name in US_STATES if name.lower() == m.group(0).lower()))
    for m in STATE_POSTAL_RE.finditer(text):
        code = m.group(0)
        # skip high-ambiguity codes unless they look like a state in an address context
        if code in _HIGH_AMBIGUITY_CODES:
            # only accept in address context
            start = m.start()
            window_before = text[max(0, start - 3):start]
            window_after = text[m.end():m.end() + 8]
            in_address_ctx = (
                re.search(r",\s*$", window_before) and re.search(r"^\s*\d{5,6}", window_after)
            )
            if not in_address_ctx:
                continue  # looks like a label, not a state
        results.append(code)
    unique_results = list(dict.fromkeys(results))
    # Normalize to postal codes and dedup to prevent returning both 'New York' and 'NY'
    _postal_seen: set = set()
    _normalized: list = []
    for r in unique_results:
        # If r is a full name, convert to postal code
        code = US_STATES.get(r) or US_STATES.get(r.title()) or r
        if code not in _postal_seen:
            _postal_seen.add(code)
            _normalized.append(code)
    unique_results = _normalized
    return unique_results if unique_only else ", ".join(unique_results)

def detect_gender_values(text: str) -> List[str]:
    """Return normalized gender mentions. Uses word boundaries to avoid false matches inside words like 'transaction'."""
    if not text:
        return []
    found = set()
    if re.search(r"\bfemale\b", text, re.IGNORECASE):
        found.add("female")
    # Use a negative lookbehind so 'male' doesn't match inside 'female'
    if re.search(r"(?<![A-Za-z])(?<!fe)male\b", text, re.IGNORECASE):
        found.add("male")
    if re.search(r"\bnon[\s\-]?binary\b", text, re.IGNORECASE):
        found.add("non-binary")
    # 'transgender' as a whole word; bare 'trans' only as a standalone token
    if re.search(r"\btransgender\b", text, re.IGNORECASE):
        found.add("transgender")
    elif re.search(r"\btrans\b", text, re.IGNORECASE):
        # don't fire if 'trans' is a prefix of a non-gender word
        if not re.search(r"\btrans(?:action|fer|mit|port|form|lat|plant|parent)", text, re.IGNORECASE):
            found.add("transgender")
    # M/F shorthand
    if re.search(r"\bm\s*[/\-]\s*f\b", text, re.IGNORECASE):
        found.add("m/f")
    # Single letter M/F only when a gender/sex keyword is present
    if re.search(r"\b(?:gender|sex)\s*[:\-]?\s*m\b", text, re.IGNORECASE):
        found.add("male")
    if re.search(r"\b(?:gender|sex)\s*[:\-]?\s*f\b", text, re.IGNORECASE):
        found.add("female")
    return list(found)

def get_confidence(pi_type: str, match_text: str, text_context: str) -> int:
    """Return 0-100 confidence. Below 40 is dropped by clean_pi_data. Account Numbers without banking context score 15; bare 9-digit passports score 45."""
    ctx = text_context or ""
    # Always high confidence
    if pi_type in ("email", "SSN Number"):
        return 100
    if pi_type == "PAN Card":
        return 98
    if pi_type == "Aadhaar Number":
        return 97
    if pi_type == "IFSC Code":
        return 96
    if pi_type == "IP Address":
        return 95
    if pi_type == "Credit/Debit Card":
        return 92
    if pi_type == "Indian Phone":
        return 92
    if pi_type == "Phone":
        return 88
    if pi_type == "Voter ID":
        return 88 if _VOTER_ID_CTX_RE.search(ctx) else 72
    # Account Number — must have a banking keyword or it gets dropped
    if pi_type == "Account Number":
        if not _has_account_context(ctx):
            return 15   
        return 90 if len(re.sub(r"\D","",match_text)) >= 12 else 80
    # Passport — letter+8digits is strong; bare 9 digits needs a keyword nearby
    if pi_type == "US Passport Number":
        if PASSPORT_KEYWORD.search(ctx):
            return 95
        if re.match(r"^[A-Z]\d{8}$", match_text.strip()):
            # letter+8digits without a passport keyword gets low confidence
            return 30  
        return 45   # bare 9 digits, no keyword — likely false positive
    if pi_type.startswith("Driving License"):
        return 95
    if pi_type == "Person Name":
        return 95 if get_nlp() is not None else 85
    if pi_type == "Age":
        return 90
    if pi_type == "Gender":
        return 85
    if pi_type == "Credit/Debit Card":
        return 92
    if pi_type == "Address":
        return 80
    if pi_type == "MRN":
        return 90 if _has_mrn_context(ctx) else 35
    if pi_type == "Insurance ID":
        return 88
    if pi_type == "Date of Birth":
        return 90
    if pi_type == "Company ID":
        return 90
    if pi_type == "State":
        return 70
    return 80
# DOB patterns
ISO_YYYY_MM_DD = r"""
\b
(?:
    (?:19|20)\d{2}                # year
    [/\-\.]
    (?:0?[1-9]|1[0-2])            # month
    [/\-\.]
    (?:0?[1-9]|[12][0-9]|3[01])   # day
)
\b
"""

# YYYY/DD/MM  year-first, day before month (less common but seen in some locales)
ISO_YYYY_DD_MM = r"""
\b
(?:
    (?:19|20)\d{2}                # year
    [/\-\.]
    (?:0?[1-9]|[12][0-9]|3[01])  # day
    [/\-\.]
    (?:0?[1-9]|1[0-2])            # month
)
\b
"""

NUMERIC_MM_DD_YYYY = r"""
\b
(?:
    (?:0?[1-9]|1[0-2])           # month
    [/\-]
    (?:0?[1-9]|[12][0-9]|3[01])  # day
    [/\-]
    (?:19|20)\d{2}               # year
)
\b
"""

MONTH_SHORT_WITH_ORDINAL = r"""
\b
(?:
    (?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)
    \s+
    (?:0?[1-9]|[12][0-9]|3[01])(?:st|nd|rd|th)?
    \s+
    (?:19|20)\d{2}
)
\b
"""

MONTH_FULL_WITH_ORDINAL = r"""
\b
(?:
    (?:January|February|March|April|May|June|July|August|September|October|November|December)
    \s+
    (?:0?[1-9]|[12][0-9]|3[01])(?:st|nd|rd|th)?
    \s+
    (?:19|20)\d{2}
)
\b
"""

ORDINAL_DAY_SHORT_MONTH = r"""
\b
(?:
    (?:0?[1-9]|[12][0-9]|3[01])(?:st|nd|rd|th)?
    \s+
    (?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)
    \s+
    (?:19|20)\d{2}
)
\b
"""

ORDINAL_DAY_FULL_MONTH = r"""
\b
(?:
    (?:0?[1-9]|[12][0-9]|3[01])(?:st|nd|rd|th)?
    \s+
    (?:January|February|March|April|May|June|July|August|September|October|November|December)
    \s+
    (?:19|20)\d{2}
)
\b
"""

# DD/MM/YYYY  Indian / European date format (day first)
NUMERIC_DD_MM_YYYY = r"""
\b
(?:
    (?:0?[1-9]|[12][0-9]|3[01])  # day (1-31)
    [/\-]
    (?:0?[1-9]|1[0-2])            # month (1-12)
    [/\-]
    (?:19|20)\d{2}               # year
)
\b
"""

DOB_PATTERNS = [
    re.compile(ISO_YYYY_MM_DD,           flags=re.IGNORECASE | re.VERBOSE),  # YYYY-MM-DD (ISO)
    re.compile(ISO_YYYY_DD_MM,           flags=re.IGNORECASE | re.VERBOSE),  # YYYY-DD-MM
    re.compile(NUMERIC_DD_MM_YYYY,       flags=re.IGNORECASE | re.VERBOSE),  # DD/MM/YYYY Indian
    re.compile(NUMERIC_MM_DD_YYYY,       flags=re.IGNORECASE | re.VERBOSE),  # MM/DD/YYYY US
    re.compile(MONTH_SHORT_WITH_ORDINAL, flags=re.IGNORECASE | re.VERBOSE),
    re.compile(MONTH_FULL_WITH_ORDINAL,  flags=re.IGNORECASE | re.VERBOSE),
    re.compile(ORDINAL_DAY_SHORT_MONTH,  flags=re.IGNORECASE | re.VERBOSE),
    re.compile(ORDINAL_DAY_FULL_MONTH,   flags=re.IGNORECASE | re.VERBOSE),
]


# Standalone PII detector — used by DOCX, PDF, and table rows
def _detect_standalone_pii(
    text: str,
    file_name: str,
    page_sheet: str,
    cell: str,
    records: List[Dict[str, Any]],
    seen: set,
) -> None:
    """Detect all PII types except name pairs (handled by emit_grouped_records). Each hit is added as an individual row."""
    if not text or not text.strip():
        return

    def _add(pi_type: str, value: str, confidence: int) -> None:
        value = (value or "").strip()
        if not value or confidence < 40:
            return
        _add_record(records, seen,
            file_name=file_name, page_sheet=page_sheet,
            cell=cell, pi_type=pi_type,
            value=value, confidence=confidence)

    # SSN
    for v in find_ssn_in_text(text):
        _add("SSN Number", v, 100)

    # Voter ID (always scan regardless of context)
    for m in re.finditer(pi_PATTERNS["Voter ID"], text):
        _add("Voter ID", m.group(0), 90)

    # Standalone person names
    names = extract_person_names(text)
    for name in names:
        # Only accept title-case names
        if re.match(r"^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+$", name):
            _add("Person Name", name, 90)

    # Pre-collect Aadhaar spans to avoid Phone double-matching them
    _aadhaar_spans = [(m.start(), m.end()) for m in re.finditer(pi_PATTERNS["Aadhaar Number"], text)]
    def _overlaps_aadhaar(s, e):
        return any(not (e <= a or s >= b) for a, b in _aadhaar_spans)

    # Phone — grab opening bracket if the regex missed it
    for m in re.finditer(pi_PATTERNS["Phone"], text):
        v = m.group(0)
        s = m.start()
        if s > 0 and text[s-1] == "(" and not v.startswith("("):
            v = "(" + v
        # skip SSN format
        if re.match(r"^\d{3}-\d{2}-\d{4}$", v.strip()):
            continue
        # skip if overlapping an Aadhaar number
        if _overlaps_aadhaar(m.start(), m.end()):
            continue
        # skip if it looks like an IP address fragment
        if '.' in v and re.search(r'\d+\.\d+\.\d+', v):
            continue
        _add("Phone", v, 88)

    # Indian Phone — skip if the number looks like a bank account number
    _ACCT_NEARBY_RE = re.compile(r"\b(?:account|acct|a\/c|bank\s+account)\b", re.IGNORECASE)
    for m in re.finditer(pi_PATTERNS["Indian Phone"], text):
        v = m.group(0)
        # check 40-char window before the match for account keywords
        window = text[max(0, m.start()-40):m.start()]
        if _ACCT_NEARBY_RE.search(window):
            continue  # this is an account number, not a phone
        _add("Indian Phone", m.group(0), 92)

    # Email (standalone — not paired here)
    for m in re.finditer(pi_PATTERNS["email"], text):
        _add("email", m.group(0), 100)

    # Credit/Debit Card — Luhn validated
    for m in re.finditer(pi_PATTERNS["Credit/Debit Card"], text):
        v = m.group(0)
        if passes_luhn(v):
            _add("Credit/Debit Card", v, 92)

    # IP Address
    for m in re.finditer(pi_PATTERNS["IP Address"], text):
        _add("IP Address", m.group(0), 95)

    # Aadhaar
    for m in re.finditer(pi_PATTERNS["Aadhaar Number"], text):
        _add("Aadhaar Number", m.group(0), 97)

    # PAN Card
    for m in re.finditer(pi_PATTERNS["PAN Card"], text):
        _add("PAN Card", m.group(0), 98)

    # IFSC Code
    for m in re.finditer(pi_PATTERNS["IFSC Code"], text):
        _add("IFSC Code", m.group(0), 96)

    # US Passport — requires keyword nearby and specifically before the value
    for m in re.finditer(pi_PATTERNS["US Passport Number"], text):
        v = m.group(0).strip()
        # keyword must be within 25 chars before this value
        window_before = text[max(0, m.start()-25): m.start()]
        kw_nearby = bool(PASSPORT_KEYWORD.search(window_before))
        if re.match(r"^[A-Z]\d{8}$", v):
            if kw_nearby or PASSPORT_KEYWORD.search(text):
                _add("US Passport Number", v, 95)
        elif PASSPORT_9D_RE.match(v) and kw_nearby:
            # 9-digit: only emit when passport keyword is directly before this value
            _add("US Passport Number", v, 95)

    # Driving license (context-gated)
    for state, full_value, _ in iter_dl_matches(text, col_name=None, require_context=True):
        _add(f"Driving License ({state})", full_value, 95)

    # Address — collect all candidates then drop substrings of longer matches
    flat = re.sub(r"[\r\n\t]+", " ", text)
    _raw_addrs = []
    for m in re.findall(ADDRESS_PATTERN, flat, flags=re.IGNORECASE):
        v = m.strip()
        if v:
            _raw_addrs.append((v, 80))
    for m in ADDRESS_SUITE_PATTERN.finditer(flat):
        v = m.group(0).strip()
        if v and len(v) > 8:
            _raw_addrs.append((v, 78))
    _raw_addrs_sorted = sorted(_raw_addrs, key=lambda x: len(x[0]), reverse=True)
    _kept_addrs = []
    for v, conf in _raw_addrs_sorted:
        if not any(v != longer and v in longer for longer, _ in _kept_addrs):
            _kept_addrs.append((v, conf))
    for v, conf in _kept_addrs:
        _add("Address", v, conf)

    # City/state/ZIP fragment even without a street number
    _CITY_STATE_PIN_RE = re.compile(
        r"\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?),\s*([A-Z]{2})\s+(\d{5,6})\b"
    )
    for m in _CITY_STATE_PIN_RE.finditer(flat):
        frag = m.group(0).strip()
        # Only add if not already covered by a longer address
        if not any(frag in longer for longer, _ in _kept_addrs):
            _add("Address", frag, 78)

    # Standalone date of birth
    _dob_iso = first_valid_dob_iso(text)
    if _dob_iso:
        _add("Date of Birth", _dob_iso, 90)

    # Account Number — banking keyword required; skip Luhn and Indian phone matches
    _indian_phone_re = re.compile(r"^[6-9]\d{9}$")
    # Build set of passport values that were actually emitted, so we don't skip routing numbers
    _passport_vals = set()
    for _m in re.finditer(pi_PATTERNS["US Passport Number"], text):
        _v = _m.group(0).strip()
        _wb = text[max(0, _m.start()-25): _m.start()]
        if PASSPORT_KEYWORD.search(_wb) or (re.match(r"^[A-Z]\d{8}$", _v) and PASSPORT_KEYWORD.search(text)):
            _passport_vals.add(_v)
    for m in re.finditer(pi_PATTERNS["Account Number"], text):
        v = m.group(0)
        digits = re.sub(r"\D", "", v)
        if passes_luhn(v):
            continue  # already flagged as a credit card
        # skip if already identified as a passport number
        if v.strip() in _passport_vals:
            continue
        # override if an account keyword is right before the value
        explicit = _is_explicit_account(v, text)
        if _indian_phone_re.match(digits) and not explicit:
            continue
        if not _has_account_context(text):
            continue
        # tag as routing number if routing keyword is nearby
        if _is_routing_number(v, text):
            _add("Routing Number", v, 90)
        else:
            _add("Account Number", v, 90 if len(digits) >= 12 else 80)

    # MRN — medical keyword required
    for m in MRN_RE.finditer(text):
        if _has_mrn_context(text):
            _add("MRN", m.group(1) or m.group(2) or "", 90)

    # Insurance ID
    for m in INSURANCE_ID_RE.finditer(text):
        _add("Insurance ID", m.group(1) or m.group(2) or "", 88)

    # Company ID — EMP-prefixed values are unambiguous, others need positive context
    for m in ALNUM_COMPANY_ID_RE.finditer(text):
        token = m.group(0)
        if is_company_id(token, context=text):
            _add("Company ID", token, 90)

    # Gender
    for val in detect_gender_values(text):
        _add("Gender", val, 85)


# PDF scanner

def _normalize_pdf_page(text: str) -> str:
    """Fix PyPDF2 extraction artifacts before PII scanning (broken numbers, split names, etc.)."""
    import re as _r

    text = _r.sub(r"(\w) +-", r"\1-", text)

    for _ in range(6):
        lines = text.split("\n")
        out = []
        i = 0
        changed = False
        while i < len(lines):
            cur = lines[i].rstrip()
            if i + 1 < len(lines):
                nxt = lines[i + 1]
                # join lines ending with a hyphen (split words/numbers)
                if cur.endswith("-"):
                    out.append(cur + nxt.lstrip())
                    i += 2
                    changed = True
                    continue
                # Join lines ending with comma+space where next continues an address
                # e.g. "509 River Rd, Suite 755, " + "Denver, CO 98706"
                elif cur.endswith(", ") or cur.endswith(","):
                    nxt_stripped = nxt.strip()
                    # only join if next line starts with a capital word (address continuation)
                    if nxt_stripped and nxt_stripped[0].isupper() and not _r.match(r"^(Scenario|Section|Page|Note|Dear|Hello|From|To|Subject)\b", nxt_stripped):
                        out.append(cur.rstrip() + " " + nxt_stripped)
                        i += 2
                        changed = True
                        continue
            out.append(cur)
            i += 1
        text = "\n".join(out)
        if not changed:
            break

    text = _r.sub(r"\((\d{3})\)\s*\n\s*(\d)", r"(\1) \2", text)

    # must run before the name-join pass
    _CITY_TOKENS_P4 = {
        "MIAMI","SEATTLE","RALEIGH","ATLANTA","DALLAS","DENVER","DETROIT","CHICAGO",
        "HOUSTON","PHOENIX","BOSTON","AUSTIN","PORTLAND","MEMPHIS","NASHVILLE",
        "ANGELES","FRANCISCO","ANTONIO","DIEGO","JOSE","YORK","ORLEANS","PASO",
        "WORTH","SPRINGS","VISTA","RAPIDS","ROCK","SALEM","ANA","LAUDERDALE",
        "VEGAS","ARLINGTON","TAMPA","ORLANDO","INDIANAPOLIS","CHARLOTTE","COLUMBUS",
        "SACRAMENTO","FRESNO","MESA","OMAHA","CLEVELAND","MINNEAPOLIS","WICHITA",
        "NEWARK","PITTSBURGH","GREENSBORO","LINCOLN","CINCINNATI","TOLEDO","JERSEY",
        "PLANO","IRVINE","LAREDO","MADISON","DURHAM","LUBBOCK","GARLAND","NORFOLK",
        "SCOTTSDALE","CHANDLER","ROUGE","HIALEAH","FREMONT","RICHMOND","BAKERSFIELD",
        "GILBERT","BIRMINGHAM","ROCHESTER","SPOKANE","MONTGOMERY","GLENDALE","TACOMA",
        "AKRON","SHREVEPORT","TALLAHASSEE","HUNTSVILLE","WORCESTER","KNOXVILLE",
        "PROVIDENCE","BROWNSVILLE","BUFFALO","ANCHORAGE","HONOLULU","ANAHEIM","AURORA",
        "COLUMBIA","AUGUSTA",
    }
    _CITY_NAME_DATA_RE = _r.compile(
        r"^([A-Z][a-z]+)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\s+[\d\(\[].+)$"
    )
    _CITY_NAME_ONLY_RE = _r.compile(
        r"^([A-Z][a-z]+)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)$"
    )
    pre_split_lines = text.split("\n")
    pre_split_out = []
    for ln in pre_split_lines:
        stripped = ln.strip()
        m_a = _CITY_NAME_DATA_RE.match(stripped)
        if m_a and m_a.group(1).upper() in _CITY_TOKENS_P4:
            pre_split_out.append(m_a.group(1))
            pre_split_out.append(m_a.group(2))
            continue
        m_b = _CITY_NAME_ONLY_RE.match(stripped)
        if m_b and m_b.group(1).upper() in _CITY_TOKENS_P4:
            pre_split_out.append(m_b.group(1))
            pre_split_out.append(m_b.group(2))
            continue
        pre_split_out.append(ln)
    text = "\n".join(pre_split_out)


    lines = text.split("\n")
    out = []
    i = 0
    while i < len(lines):
        cur = lines[i].strip()
        if _r.match(r"^[A-Z][a-z]+$", cur) and i + 1 < len(lines):
            # Do NOT join if this lone word is a city token (prevents Francisco+Isabella)
            if cur.upper() in _CITY_TOKENS_P4:
                out.append(lines[i])
                i += 1
                continue
            nxt = lines[i + 1].strip()
            # Join if next starts with CapWord followed by digit/bracket (table data)
            if _r.match(r"^[A-Z][a-z]+\s+[\d\(\[]", nxt):
                out.append(cur + " " + nxt)
                i += 2
                continue

            if _r.match(r"^[A-Z][a-z]+$", nxt) and nxt.upper() not in _CITY_TOKENS_P4:
                out.append(cur + " " + nxt)
                i += 2
                continue
        out.append(lines[i])
        i += 1

    return "\n".join(out)


def find_pi_in_pdf(pdf_file) -> List[Dict[str, Any]]:

    records: List[Dict[str, Any]] = []
    seen: set[tuple] = set()  # for _add_record de-dup (per file)
    file_name = Path(getattr(pdf_file, "name", "document.pdf")).name

    # safe reset for file-like streams
    try:
        pdf_file.seek(0)
    except Exception:
        pass

    try:
        reader = PdfReader(pdf_file)
    except Exception as e:
        import traceback; traceback.print_exc()
        print(f"✖ PdfReader failed for {file_name}: {e}")
        return records

    for page_idx, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""

        if not page_text.strip():
            continue

        # normalize the page
        page_text = _normalize_pdf_page(page_text)

 
        raw_blocks = re.split(r"\n\s*\n", page_text)
        blocks: List[str] = []
        for b in raw_blocks:
            b = (b or "").strip()
            if not b:
                continue
            parts = [p.strip() for p in b.split("\n") if p.strip()]
            if len(parts) > 3:  
                blocks.extend(parts)
            else:
                blocks.append(b)


        for bt in blocks:
            grouped = emit_grouped_records(
                file_name=file_name,
                page_or_sheet=f"Page {page_idx}",
                cell="",
                block_text=bt,
                pair_only=True,
            )
            for rec in grouped:
                _add_record(
                    records, seen,
                    file_name=rec["File"],
                    page_sheet=rec["Page/Sheet"],
                    cell=rec["Cell"],
                    pi_type=rec["PI Type"],
                    value=rec["Detected Value"],
                    confidence=rec["Confidence (%)"],
                )
            _detect_standalone_pii(bt, file_name, f"Page {page_idx}", "", records, seen)

    try:
        pdf_file.seek(0)
    except Exception:
        pass

    try:
        with pdfplumber.open(pdf_file) as pdf:
            for page_idx, page in enumerate(pdf.pages, start=1):
                tables = page.extract_tables() or []
                for tbl_idx, table in enumerate(tables, start=1):
                    if not table:
                        continue

                    all_row_lines: List[str] = []
                    for r_idx, row in enumerate(table, start=1):
                        if not row:
                            continue
                        line = " | ".join(_to_text(c) for c in row if _to_text(c))
                        if not line:
                            continue
                        all_row_lines.append(line)

                        grouped = emit_grouped_records(
                            file_name=file_name,
                            page_or_sheet=f"Page {page_idx} (Table {tbl_idx})",
                            cell=f"Row {r_idx}",
                            block_text=line,
                            pair_only=True,
                        )
                        for rec in grouped:
                            _add_record(
                                records, seen,
                                file_name=rec["File"],
                                page_sheet=rec["Page/Sheet"],
                                cell=rec["Cell"],
                                pi_type=rec["PI Type"],
                                value=rec["Detected Value"],
                                confidence=rec["Confidence (%)"],
                            )
                        # standalone PII for this table row
                        _detect_standalone_pii(line, file_name,
                            f"Page {page_idx} (Table {tbl_idx})", f"Row {r_idx}", records, seen)

                    # whole-table combined block is intentionally disabled to prevent cross-person pairings
    except Exception as e:
        import traceback; traceback.print_exc()
        print(f"✖ pdfplumber tables failed for {file_name}: {e}")

    # Cross-block pairing using a sliding window per page
    try:
        pdf_file.seek(0)
    except Exception:
        pass
    try:
        reader2 = PdfReader(pdf_file)
    except Exception:
        reader2 = None

    _SCENARIO_BREAK_PDF = re.compile(
        r"^(?:scenario\s+\d|section\s+\d|page\s+\d|scenario\s+[·\-]|note:|disclaimer)", re.IGNORECASE
    )
    WINDOW_PDF = 8

    if reader2:
        for page_idx, page in enumerate(reader2.pages, start=1):
            try:
                page_text = page.extract_text() or ""
            except Exception:
                page_text = ""
            if not page_text.strip():
                continue
            page_text = _normalize_pdf_page(page_text)

            # Collect all non-empty lines for this page
            page_lines: List[str] = [ln.strip() for ln in page_text.split("\n") if ln.strip()]

            for i, anchor_line in enumerate(page_lines):
                anchor_names = extract_person_names(anchor_line)
                if not anchor_names:
                    continue

                # If the anchor already has an email, pair within that line only to avoid stealing from adjacent rows.
                _anchor_has_email = bool(re.search(pi_PATTERNS["email"], anchor_line))
                if _anchor_has_email:
                    for rec in emit_grouped_records(
                        file_name, f"Page {page_idx}", "", anchor_line, pair_only=True
                    ):
                        _add_record(
                            records, seen,
                            file_name=rec["File"], page_sheet=rec["Page/Sheet"],
                            cell=rec["Cell"], pi_type=rec["PI Type"],
                            value=rec["Detected Value"], confidence=rec["Confidence (%)"],
                        )
                    continue   # don't extend window for this anchor

                # Build a combined window — stop if a line has its own name+email pair
                window_parts: List[str] = [anchor_line]
                for j in range(i + 1, min(i + WINDOW_PDF, len(page_lines))):
                    nxt = page_lines[j]
                    if _SCENARIO_BREAK_PDF.match(nxt):
                        break
                    # If this next line contains its own name+email pair, stop here
                    _nxt_has_name = bool(extract_person_names(nxt))
                    _nxt_has_email = bool(re.search(pi_PATTERNS["email"], nxt))
                    if _nxt_has_name and _nxt_has_email:
                        break
                    window_parts.append(nxt)
                if len(window_parts) <= 1:
                    continue
                combined = " ".join(window_parts)
                # emit pairs from combined window
                for rec in emit_grouped_records(
                    file_name, f"Page {page_idx}", "", combined, pair_only=True
                ):
                    _add_record(
                        records, seen,
                        file_name=rec["File"], page_sheet=rec["Page/Sheet"],
                        cell=rec["Cell"], pi_type=rec["PI Type"],
                        value=rec["Detected Value"], confidence=rec["Confidence (%)"],
                    )

    # No other detectors here — this keeps output minimal and paired.
    return records


# Excel scanner

def find_pi_in_excel(excel_file) -> List[Dict[str, Any]]:

    records: List[Dict[str, Any]] = []
    seen: set[tuple] = set()
    file_name = Path(getattr(excel_file, "name", "workbook.xlsx")).name

    # stream reset if applicable
    try:
        excel_file.seek(0)
    except Exception:
        pass

    xls = pd.ExcelFile(excel_file)
    use_spacy = get_nlp() is not None  # compute once per function

    for sheet_name in xls.sheet_names:
        df = xls.parse(sheet_name=sheet_name, dtype=str)

        # header hints
        # name_hint_cols: all name-like columns (FullName, FirstName, LastName)
        name_hint_cols = {c for c in df.columns if NAME_COL_HINT_RE.search(str(c))}
        # full_name_cols: only FullName/Name columns, preferred for pair detection
        full_name_cols = {c for c in df.columns if FULL_NAME_COL_RE.search(str(c))}
        id_hint_cols   = {c for c in df.columns if ID_COL_HINT_RE.search(str(c))}
        dob_hint_cols  = {c for c in df.columns if DOB_COL_HINT_RE.search(str(c))}

        for row_idx, row in df.iterrows():
            # First pass: collect Name / Company ID / DOB / Email hits
            row_name_hits:    List[tuple[str, str]] = []
            row_id_hits:      List[tuple[str, str]] = []
            row_dob_hits:     List[tuple[str, str]] = []
            row_email_hits:   List[tuple[str, str]] = []
            row_phone_hits:   List[tuple[str, str]] = []
            row_addr_hits:    List[tuple[str, str]] = []
            row_state_hits:   List[tuple[str, str]] = []
            row_gender_hits:  List[tuple[str, str]] = []
            row_city_hits:    List[tuple[str, str]] = []

            for col_name, cell_value in row.items():
                text = _to_text(cell_value)
                if not text:
                    continue

                col_is_name_hint = NAME_COL_HINT_RE.search(str(col_name))
                col_is_full_name = FULL_NAME_COL_RE.search(str(col_name))
                words = text.split()
                if col_is_name_hint:
                    if not re.search(r'\d', text) and 1 <= len(words) <= 5:
                        row_name_hits.append((text.strip(), str(col_name)))
                else:
                    names_here = extract_person_names(text)
                    if names_here:
                        if len(words) <= 4 and not re.search(r'\d', text):
                            row_name_hits.append((text.strip(), str(col_name)))
                        else:
                            row_name_hits.append((names_here[0], str(col_name)))

                # Email
                email_m = re.findall(pi_PATTERNS["email"], text)
                for em in email_m:
                    row_email_hits.append((em, str(col_name)))

                # Phone (only from phone-hinted columns to avoid false positives)
                _col_lo_p = str(col_name).lower()
                if any(k in _col_lo_p for k in ('phone', 'mobile', 'cell', 'contact', 'tel')):
                    for pm in re.finditer(pi_PATTERNS["Phone"], text):
                        row_phone_hits.append((pm.group(0), str(col_name)))
                    for pm in re.finditer(pi_PATTERNS["Indian Phone"], text):
                        row_phone_hits.append((pm.group(0), str(col_name)))

                # Address (only from address-hinted columns)
                if any(k in _col_lo_p for k in ('address', 'street', 'addr')):
                    if text.strip():
                        row_addr_hits.append((text.strip(), str(col_name)))

                # State (only from state-hinted columns)
                if any(k in _col_lo_p for k in ('state', 'province', 'region')):
                    if text.strip():
                        row_state_hits.append((text.strip(), str(col_name)))

                # City (only from city-hinted columns)
                if any(k in _col_lo_p for k in ('city', 'town', 'district')):
                    if text.strip() and re.match(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*$', text.strip()):
                        row_city_hits.append((text.strip(), str(col_name)))

                # Gender (only from gender-hinted columns)
                if any(k in _col_lo_p for k in ('gender', 'sex')):
                    gvals = detect_gender_values(text)
                    for gv in gvals:
                        row_gender_hits.append((gv, str(col_name)))

                # Company ID
                _col_lo = str(col_name).lower()
                _skip_cid = (
                    any(k in _col_lo for k in ('passport','mrn','medical','routing','route',
                                               'insurance','insurance_id','aba','sort_code'))
                    or re.match(r'^[A-Z]\d{8}$', text)
                    or re.match(r'^MRN', text, re.IGNORECASE)
                    or re.match(r'^INS', text, re.IGNORECASE)
                    or re.match(r'^[A-Z]{3}\d{7}$', text)
                )
                if not _skip_cid:
                    if ALNUM_COMPANY_ID_RE.fullmatch(text) or COMPANY_ID_RE_NUM.fullmatch(text):
                        row_id_hits.append((text, str(col_name)))

                # DOB
                dob_iso = first_valid_dob_iso(text)
                if dob_iso:
                    row_dob_hits.append((dob_iso, str(col_name)))

            # prefer full-name columns for pair detection; fall back to any name hit
            def pick_best(hits: List[tuple[str, str]], preferred_cols: set[str]) -> tuple[str | None, str | None]:
                if not hits:
                    return None, None
                for val, col in hits:
                    if col in preferred_cols:
                        return val, col
                return hits[0]


            email_hint_cols  = {c for c in df.columns if re.search(r'\b(e[\-_]?mail|email[\s_]?address)\b', str(c), re.IGNORECASE)}
            phone_hint_cols  = {c for c in df.columns if re.search(r'\b(phone|mobile|cell|tel)\b', str(c), re.IGNORECASE)}
            addr_hint_cols   = {c for c in df.columns if re.search(r'\b(address|street|addr)\b', str(c), re.IGNORECASE)}
            state_hint_cols  = {c for c in df.columns if re.search(r'\b(state|province)\b', str(c), re.IGNORECASE)}
            city_hint_cols   = {c for c in df.columns if CITY_COL_HINT_RE.search(str(c))}
            gender_hint_cols = {c for c in df.columns if re.search(r'\b(gender|sex)\b', str(c), re.IGNORECASE)}

            name_val,   name_col   = pick_best(row_name_hits,   full_name_cols or name_hint_cols)
            id_val,     id_col     = pick_best(row_id_hits,     id_hint_cols)
            dob_val,    dob_col    = pick_best(row_dob_hits,    dob_hint_cols)
            email_val,  email_col  = pick_best(row_email_hits,  email_hint_cols)
            phone_val,  phone_col  = pick_best(row_phone_hits,  phone_hint_cols)
            addr_val,   addr_col   = pick_best(row_addr_hits,   addr_hint_cols)
            state_val,  state_col  = pick_best(row_state_hits,  state_hint_cols)
            city_val,   city_col   = pick_best(row_city_hits,   city_hint_cols)
            gender_val, gender_col = pick_best(row_gender_hits, gender_hint_cols)

            paired_cols: set[str] = set()

            def _emit_pair(pi_type, name, val, col_a, col_b, conf=95):
                records.append({"File": file_name, "Page/Sheet": sheet_name,
                                 "Cell": f"Row {row_idx + 2}", "PI Type": pi_type,
                                 "Detected Value": f"{name}, {val}", "Confidence (%)": conf})
                paired_cols.update({col_a, col_b})

            if name_val and email_val:
                _emit_pair("Name, email",        name_val, email_val,  name_col, email_col,  98 if use_spacy else 95)
            if name_val and id_val:
                _emit_pair("Name, Company ID",   name_val, id_val,     name_col, id_col,     98 if use_spacy else 95)
            if name_val and dob_val:
                _emit_pair("Name, Date of Birth",name_val, dob_val,    name_col, dob_col,    98 if use_spacy else 95)
            if name_val and phone_val:
                _emit_pair("Name, Phone",        name_val, phone_val,  name_col, phone_col,  95)
            if name_val and addr_val:
                _emit_pair("Name, Address",      name_val, addr_val,   name_col, addr_col,   95)
            if name_val and city_val:
                _emit_pair("Name, City",         name_val, city_val,   name_col, city_col,   85)
            if name_val and state_val:
                _emit_pair("Name, State",        name_val, state_val,  name_col, state_col,  90)
            if name_val and gender_val:
                _emit_pair("Name, Gender",       name_val, gender_val, name_col, gender_col, 90)

            # Second pass: per-cell PII, skipping cells already used in a pair
            for col_name, cell_value in row.items():
                text = _to_text(cell_value)
                if not text:
                    continue

                # skip if this cell was part of a pair
                in_paired_cell = str(col_name) in paired_cols

                # Column header hints
                col_hint_is_id = bool(ID_COL_HINT_RE.search(str(col_name)))

                # skip DL for paired cells, ID-like headers, or full-match company IDs
                skip_dl_for_cell = (
                    in_paired_cell
                    or col_hint_is_id
                    or bool(ALNUM_COMPANY_ID_RE.fullmatch(text))
                    or bool(COMPANY_ID_RE_NUM.fullmatch(text))
                )

                # Bank Account column override
                col_hint_is_bank_acct = bool(BANK_ACCT_COL_HINT_RE.search(str(col_name)))
                if col_hint_is_bank_acct and re.match(r"^\d{9,18}$", text.strip()):
                    records.append({
                        "File": file_name, "Page/Sheet": sheet_name,
                        "Cell": f"{col_name} (Row {row_idx + 2})",
                        "PI Type": "Account Number", "Detected Value": text.strip(),
                        "Confidence (%)": 95,
                    })

                # Company ID override (only if not a paired cell)
                _col_lo2 = str(col_name).lower()
                _skip_cid2 = (
                    any(k in _col_lo2 for k in ('passport','mrn','medical','routing','route',
                                                'insurance','aba','sort_code'))
                    or re.match(r'^[A-Z]\d{8}$', text)
                    or re.match(r'^MRN', text, re.IGNORECASE)
                    or re.match(r'^INS', text, re.IGNORECASE)
                    or re.match(r'^[A-Z]{3}\d{7}$', text)
                )
                if not in_paired_cell and not _skip_cid2 and col_hint_is_id and (
                    ALNUM_COMPANY_ID_RE.fullmatch(text) or COMPANY_ID_RE_NUM.fullmatch(text)
                ):
                    records.append({
                        "File": file_name,
                        "Page/Sheet": sheet_name,
                        "Cell": f"{col_name} (Row {row_idx + 2})",
                        "PI Type": "Company ID",
                        "Detected Value": text,
                        "Confidence (%)": 100,
                    })
                    # no 'continue': still allow other PII (email/phone/etc.)

                # Person Name (only if not a paired cell)
                if not in_paired_cell:
                    col_is_name_col = bool(NAME_COL_HINT_RE.search(str(col_name)))
                    if col_is_name_col:
                        # trust the column header and emit the cell value directly as a Person Name
                        words_n = text.split()
                        if 1 <= len(words_n) <= 5 and not re.search(r'\d', text) and not _is_bad_label(text):
                            records.append({
                                "File": file_name,
                                "Page/Sheet": sheet_name,
                                "Cell": f"{col_name} (Row {row_idx + 2})",
                                "PI Type": "Person Name",
                                "Detected Value": text.strip(),
                                "Confidence (%)": 95,
                            })
                    else:
                        for name in extract_person_names(text):
                            records.append({
                                "File": file_name,
                                "Page/Sheet": sheet_name,
                                "Cell": f"{col_name} (Row {row_idx + 2})",
                                "PI Type": "Person Name",
                                "Detected Value": name,
                                "Confidence (%)": 95 if use_spacy else 85,
                            })

                # States — skip address/street columns since Indian state codes appear in addresses too
                _col_is_addr = any(k in str(col_name).lower() for k in ('address','street','addr','location'))
                if not _col_is_addr:
                    states_detected = detect_states(text)
                    if states_detected:
                        records.append({
                            "File": file_name,
                            "Page/Sheet": sheet_name,
                            "Cell": f"{col_name} (Row {row_idx + 2})",
                            "PI Type": "State",
                            "Detected Value": states_detected,
                            "Confidence (%)": 75,
                        })

                # DL (context-gated)
                if not skip_dl_for_cell:
                    for state, full_value, core in iter_dl_matches(text, col_name=str(col_name), require_context=True):
                        # if the cell has an explicit 2-letter state prefix, use that rather than the pattern-matched state
                        _explicit_prefix = re.match(r'^([A-Z]{2})[-\s]', text.strip())
                        if _explicit_prefix:
                            declared_state = _explicit_prefix.group(1).upper()
                            # only override when it's a known US state
                            if declared_state in US_DL_PATTERNS or declared_state in _PURE_DIGIT_DL_STATES:
                                state = declared_state
                                full_value = text.strip()  # use the original cell value as-is
                        elif not full_value.upper().startswith(state):
                            full_value = f"{state}-{full_value}"
                        records.append({
                            "File": file_name,
                            "Page/Sheet": sheet_name,
                            "Cell": f"{col_name} (Row {row_idx + 2})",
                            "PI Type": f"Driving License ({state})",
                            "Detected Value": full_value,
                            "Confidence (%)": 95,
                        })

                # MRN
                _col_is_mrn = bool(re.search(
                    r'(?:^|_|\s)(?:mrn|medical(?:record)?(?:number|no|num)?)',
                    str(col_name), re.IGNORECASE
                ))
                if _col_is_mrn:
                    # column is explicitly MRN — emit the full cell value
                    if text.strip():
                        records.append({
                            "File": file_name,
                            "Page/Sheet": sheet_name,
                            "Cell": f"{col_name} (Row {row_idx + 2})",
                            "PI Type": "MRN",
                            "Detected Value": text.strip(),
                            "Confidence (%)": 95,
                        })
                elif _has_mrn_context(text):
                    for m in MRN_RE.finditer(text):
                        mrn_val = m.group(1) or m.group(2) or ""
                        if mrn_val:
                            records.append({
                                "File": file_name,
                                "Page/Sheet": sheet_name,
                                "Cell": f"{col_name} (Row {row_idx + 2})",
                                "PI Type": "MRN",
                                "Detected Value": mrn_val,
                                "Confidence (%)": 95,
                            })

                # Routing Number — detect when column header mentions routing
                _col_is_routing = bool(re.search(r'\b(?:routing|aba|routing\s*number)\b', str(col_name), re.IGNORECASE))
                if _col_is_routing and re.match(r'^\d{9}$', text.strip()):
                    records.append({
                        "File": file_name,
                        "Page/Sheet": sheet_name,
                        "Cell": f"{col_name} (Row {row_idx + 2})",
                        "PI Type": "Routing Number",
                        "Detected Value": text.strip(),
                        "Confidence (%)": 95,
                    })

                # Insurance ID
                _col_is_insurance = bool(re.search(r'\b(?:insurance|insur|policy)\b', str(col_name), re.IGNORECASE))
                if _col_is_insurance and text.strip():
                    records.append({
                        "File": file_name,
                        "Page/Sheet": sheet_name,
                        "Cell": f"{col_name} (Row {row_idx + 2})",
                        "PI Type": "Insurance ID",
                        "Detected Value": text.strip(),
                        "Confidence (%)": 92,
                    })
                elif re.match(r'^INS[-\s]', text.strip(), re.IGNORECASE):
                    records.append({
                        "File": file_name,
                        "Page/Sheet": sheet_name,
                        "Cell": f"{col_name} (Row {row_idx + 2})",
                        "PI Type": "Insurance ID",
                        "Detected Value": text.strip(),
                        "Confidence (%)": 90,
                    })

                # Address
                for m in re.findall(ADDRESS_PATTERN, text, flags=re.IGNORECASE):
                    v = m.strip()
                    if v:
                        records.append({
                            "File": file_name,
                            "Page/Sheet": sheet_name,
                            "Cell": f"{col_name} (Row {row_idx + 2})",
                            "PI Type": "Address",
                            "Detected Value": v,
                            "Confidence (%)": 75,
                        })

                # Date of Birth (only if not in a pair)
                if not in_paired_cell:
                    dob_iso = first_valid_dob_iso(text)
                    if dob_iso:
                        records.append({
                            "File": file_name,
                            "Page/Sheet": sheet_name,
                            "Cell": f"{col_name} (Row {row_idx + 2})",
                            "PI Type": "Date of Birth",
                            "Detected Value": dob_iso,
                            "Confidence (%)": 95,
                        })

                # Other regex PII — pre-compute Aadhaar spans to avoid Phone overlap
                _cell_aadhaar = [(m.start(), m.end()) for m in re.finditer(pi_PATTERNS["Aadhaar Number"], text)]
                def _aadhaar_overlap(s, e):
                    return any(not (e <= a or s >= b) for a, b in _cell_aadhaar)

                for pi_type, pattern in pi_PATTERNS.items():
                    for m in re.finditer(pattern, text):
                        value = m.group(0)
                        # Drop Phone matches that sit inside an Aadhaar number
                        if pi_type == "Phone" and _aadhaar_overlap(m.start(), m.end()):
                            continue
                        # US Passport: 9-digit values must NOT come from routing/account/MRN columns
                        if pi_type == "US Passport Number":
                            col_str = str(col_name).lower()
                            _is_routing_col = any(k in col_str for k in ("routing","route","aba","sort_code"))
                            _is_account_col = any(k in col_str for k in ("account","acct","bank","mrn","medical","insurance"))
                            # pure 9-digit passport needs a keyword or a passport-hinted column
                            _is_9d = bool(re.match(r'^\d{9}$', value))
                            _col_is_passport = any(k in col_str for k in ("passport","pass_no","pass_num"))
                            if _is_9d and (_is_routing_col or _is_account_col) and not _col_is_passport:
                                continue  # routing/account number, not a passport
                            if _is_9d and not _col_is_passport and not PASSPORT_KEYWORD.search(text):
                                # 9-digit value in a non-passport column with no keyword — skip
                                continue
                            # letter+8digit in a passport-hinted column
                            if _col_is_passport:
                                records.append({
                                    "File": file_name,
                                    "Page/Sheet": sheet_name,
                                    "Cell": f"{col_name} (Row {row_idx + 2})",
                                    "PI Type": pi_type,
                                    "Detected Value": value,
                                    "Confidence (%)": 95,
                                })
                                continue
                        # Account Number in Excel — treat column name as additional context
                        if pi_type == "Account Number":
                            if _has_account_context(text, col_name=str(col_name)):
                                digits_v = re.sub(r"\D", "", value)
                                # 9-digit numbers near routing keyword — tag as Routing Number
                                if _is_routing_number(value, text) or "routing" in str(col_name).lower():
                                    conf = 90
                                    records.append({
                                        "File": file_name,
                                        "Page/Sheet": sheet_name,
                                        "Cell": f"{col_name} (Row {row_idx + 2})",
                                        "PI Type": "Routing Number",
                                        "Detected Value": value,
                                        "Confidence (%)": conf,
                                    })
                                    continue
                                conf = 90 if len(digits_v) >= 12 else 80
                            else:
                                continue  # drop — no banking context anywhere
                        else:
                            conf = get_confidence(pi_type, value, text)
                        records.append({
                            "File": file_name,
                            "Page/Sheet": sheet_name,
                            "Cell": f"{col_name} (Row {row_idx + 2})",
                            "PI Type": pi_type,
                            "Detected Value": value,
                            "Confidence (%)": conf,
                        })

                # Gender
                for val in detect_gender_values(text):
                    records.append({
                        "File": file_name,
                        "Page/Sheet": sheet_name,
                        "Cell": f"{col_name} (Row {row_idx + 2})",
                        "PI Type": "Gender",
                        "Detected Value": val,
                        "Confidence (%)": 85,
                    })

                # City (column-header triggered)
                if CITY_COL_HINT_RE.search(str(col_name)) and text:
                    # Accept values that look like a proper-cased city name
                    if re.match(r"^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*$", text.strip()):
                        records.append({
                            "File": file_name,
                            "Page/Sheet": sheet_name,
                            "Cell": f"{col_name} (Row {row_idx + 2})",
                            "PI Type": "City",
                            "Detected Value": text.strip(),
                            "Confidence (%)": 80,
                        })

    # Whole-sheet pair detection for key-value layout sheets only.
    # Only runs when per-row scanning produced no pairs (vertical layout).
    # Never runs on tabular data (each row = one person) to avoid cross-row false pairings.
    _KV_SAFE_PAIRS = {"Name, email", "Name, Date of Birth", "Name, Company ID"}
    _perrow_pair_types = {r.get("PI Type","") for r in records if r.get("PI Type","").startswith("Name,")}
    _has_perrow_pairs = bool(_perrow_pair_types)
    if not _has_perrow_pairs:
        for sheet_name in xls.sheet_names:
            df_kv = xls.parse(sheet_name=sheet_name, dtype=str)
            all_cell_values: List[str] = []
            for _, row in df_kv.iterrows():
                for val in row:
                    v = _to_text(val)
                    if v:
                        all_cell_values.append(v)
            if len(all_cell_values) < 2:
                continue
            sheet_block = " | ".join(all_cell_values)
            for rec in emit_grouped_records(file_name, sheet_name, "", sheet_block, pair_only=True):
                if rec["PI Type"] not in _KV_SAFE_PAIRS:
                    continue
                _add_record(
                    records,
                    set(),
                    file_name=rec["File"],
                    page_sheet=rec["Page/Sheet"],
                    cell=rec["Cell"],
                    pi_type=rec["PI Type"],
                    value=rec["Detected Value"],
                    confidence=rec["Confidence (%)"],
                )

            # Also detect standalone account numbers in KV sheets when banking context exists
            if _has_account_context(sheet_block):
                _kv_seen_acct: set = set()
                for _m in re.finditer(pi_PATTERNS["Account Number"], sheet_block):
                    _av = _m.group(0)
                    if _av in _kv_seen_acct:
                        continue
                    _kv_seen_acct.add(_av)
                    # skip Luhn-valid values (credit cards) and Indian-phone-shaped values
                    _ind_re_kv = re.compile(r'^[6-9]\d{9}$')
                    if passes_luhn(_av):
                        continue
                    if _ind_re_kv.match(re.sub(r"\D", "", _av)) and not _is_explicit_account(_av, sheet_block):
                        continue
                    _add_record(
                        records,
                        set(),
                        file_name=file_name,
                        page_sheet=sheet_name,
                        cell="",
                        pi_type="Account Number",
                        value=_av,
                        confidence=85,
                    )

    return records

# DOCX scanner
def find_pi_in_docx(docx_file) -> List[Dict[str, Any]]:
    records: List[Dict[str, Any]] = []
    file_name = Path(getattr(docx_file, "name", "document.docx")).name
    seen: set[tuple] = set()  # for _add_record de-dup (per file)

    # load to a temp file first
    try:
        docx_file.seek(0)
    except Exception:
        pass
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        try:
            tmp.write(docx_file.read())
        except Exception:
            docx_file.seek(0); shutil.copyfileobj(docx_file, tmp)
        tmp_path = tmp.name
    try:
        document = Document(tmp_path)
    finally:
        try: os.remove(tmp_path)
        except Exception: pass

    # Paragraphs — pairs and standalone PII
    para_texts: List[tuple[int, str]] = []  # (idx, text) for cross-paragraph pairing
    for idx, para in enumerate(getattr(document, "paragraphs", []) or [], start=1):
        text = _to_text(getattr(para, "text", ""))
        if not text:
            continue
        para_texts.append((idx, text))
        grouped = emit_grouped_records(file_name, f"Paragraph {idx}", "NA", text, pair_only=True)
        for rec in grouped:
            _add_record(records, seen,
                file_name=rec["File"], page_sheet=rec["Page/Sheet"],
                cell=rec["Cell"], pi_type=rec["PI Type"],
                value=rec["Detected Value"], confidence=rec["Confidence (%)"])
        _detect_standalone_pii(text, file_name, f"Paragraph {idx}", "NA", records, seen)


    _SCENARIO_BREAK_RE = re.compile(
        r"^(?:scenario\s+\d|section\s+\d|dear\s+hr|subject:|to:|from:|page\s+\d)", re.IGNORECASE
    )
    WINDOW = 5
    for i in range(len(para_texts)):
        anchor_idx, anchor_text = para_texts[i]
        # only start a window from a paragraph that contains a name
        anchor_names = extract_person_names(anchor_text)
        if not anchor_names:
            continue
        # collect this paragraph + next WINDOW-1 paragraphs
        window_parts: List[str] = [anchor_text]
        for j in range(i + 1, min(i + WINDOW, len(para_texts))):
            _, nxt_text = para_texts[j]
            if _SCENARIO_BREAK_RE.match(nxt_text):
                break
            window_parts.append(nxt_text)
        if len(window_parts) <= 1:
            continue  # nothing to extend
        combined = " ".join(window_parts)
        # run pair detection on combined block — only new cross-paragraph pairs are emitted
        combined_pairs = emit_grouped_records(
            file_name, f"Paragraph {anchor_idx}", "NA", combined, pair_only=True
        )
        for rec in combined_pairs:
            _add_record(records, seen,
                file_name=rec["File"], page_sheet=rec["Page/Sheet"],
                cell=rec["Cell"], pi_type=rec["PI Type"],
                value=rec["Detected Value"], confidence=rec["Confidence (%)"])

    # Tables — one row per proximity block, pairs only
    for t_idx, table in enumerate(getattr(document, "tables", []) or [], start=1):
        try:
            if table is None or len(table.rows) == 0:
                continue

            # find the first non-empty row to infer column count
            first_row = next((r for r in table.rows if getattr(r, "cells", None) and len(r.cells) > 0), None)
            if first_row is None:
                continue
            n_cols = len(first_row.cells)

            # infer if first row is a header
            header_cells = table.rows[0].cells if len(table.rows[0].cells) > 0 else []
            headers = [_to_text(c.text) for c in header_cells] if header_cells else []
            has_header = any([
                (any(NAME_COL_HINT_RE.search(h) for h in headers) if headers else False),
                (any(ID_COL_HINT_RE.search(h)   for h in headers) if headers else False),
                (any(DOB_COL_HINT_RE.search(h)  for h in headers) if ('DOB_COL_HINT_RE' in globals() and headers) else False),
            ])
            start_row = 1 if has_header else 0
            cols = headers if (has_header and headers and len(headers) == n_cols) else [f"Col {i+1}" for i in range(n_cols)]

            all_row_lines_docx: List[str] = []

            for r_idx in range(start_row, len(table.rows)):
                row = table.rows[r_idx]
                cells = getattr(row, "cells", None)
                if not cells:
                    continue
                width = min(n_cols, len(cells))
                # build a single row string as the proximity unit
                line = " | ".join(_to_text(cells[c].text) for c in range(width) if _to_text(cells[c].text))
                if not line:
                    continue
                all_row_lines_docx.append(line)

                grouped = emit_grouped_records(file_name, f"Table {t_idx}", f"Row {r_idx + 1}", line, pair_only=True)
                for rec in grouped:
                    _add_record(records, seen,
                        file_name=rec["File"],
                        page_sheet=rec["Page/Sheet"],
                        cell=rec["Cell"],
                        pi_type=rec["PI Type"],
                        value=rec["Detected Value"],
                        confidence=rec["Confidence (%)"],
                    )
                # standalone PII for this table row
                _detect_standalone_pii(line, file_name, f"Table {t_idx}", f"Row {r_idx + 1}", records, seen)

            # No per-cell fallback in DOCX tables.
            # Whole-table cross-row pairing is disabled to avoid cross-person pairings in contact tables.

        except Exception as e:
            import traceback; traceback.print_exc()
            print(f"✖ DOCX table parse failed (Table {t_idx}): {e}")
            continue

    return records

# Master detector — combines per-type scanners
def detect_pi(files: List[Any]) -> pd.DataFrame:
    """Scan a list of uploaded files and return a combined DataFrame of detected PII. Supports PDF, DOCX, and XLSX/XLS."""
    supported_ext = {"pdf", "docx", "xlsx", "xls"}
    all_records: List[Dict[str, Any]] = []

    for f in files:
        name = getattr(f, "name", "")
        ext = name.lower().split(".")[-1]
        if ext not in supported_ext:
            continue

        try:
            if ext == "pdf":
                all_records.extend(find_pi_in_pdf(f))
            elif ext == "docx":
                all_records.extend(find_pi_in_docx(f))
            elif ext in {"xlsx", "xls"}:
                all_records.extend(find_pi_in_excel(f))
        except Exception as e:
            print(f"✖ Error processing {name}: {e}")

    return pd.DataFrame(all_records, columns=["File", "Page/Sheet", "Cell", "PI Type", "Detected Value", "Confidence (%)"])

# Redaction support
def load_pi_mapping(pi_df: pd.DataFrame, selected_pi_types: List[str]) -> Dict[str, List[str]]:
    """Build a {file_name: [values_to_redact]} mapping. Pair values are split into components, and name words are added individually so FirstName/LastName cells are also caught."""
    if pi_df is None or pi_df.empty:
        return {}
    required_cols = ["File", "Detected Value", "PI Type"]
    for col in required_cols:
        if col not in pi_df.columns:
            raise ValueError(f"PI DataFrame must contain '{col}' column")

    ALL_PAIR_TYPES = {"Name, email", "Name, Date of Birth", "Name, Company ID",
                      "Name, Phone", "Name, Address", "Name, City", "Name, State",
                      "Name, Gender", "Name, Email", "Name, DOB"}

    # check if any name-related pair type was selected
    redacting_names = bool(set(selected_pi_types) & ALL_PAIR_TYPES)
    # standalone Person Name also triggers name redaction
    if "Person Name" in selected_pi_types:
        redacting_names = True

    mapping: Dict[str, List[str]] = {}

    def _add(fname: str, val: str) -> None:
        val = val.strip()
        if val:
            mapping.setdefault(fname, []).append(val)

    def _add_name_with_parts(fname: str, full_name: str) -> None:
        """Add the full name and each word component so FirstName/LastName cells are also redacted."""
        full_name = full_name.strip()
        if not full_name:
            return
        _add(fname, full_name)
        parts = full_name.split()
        if len(parts) > 1:
            for part in parts:
                if len(part) >= 2:  # skip initials like "J."
                    _add(fname, part)

    df_selected = pi_df[pi_df['PI Type'].isin(selected_pi_types)].dropna(subset=["File", "Detected Value"])
    for _, row in df_selected.iterrows():
        fname   = str(row["File"]).strip()
        val     = str(row["Detected Value"]).strip()
        pi_type = str(row.get("PI Type", "")).strip()

        if pi_type in ALL_PAIR_TYPES:
            parts = [p.strip() for p in val.split(", ", 1)]
            # first part is always the name
            if parts:
                _add_name_with_parts(fname, parts[0])
            # second part (email/DOB/ID) — add as-is
            if len(parts) > 1:
                _add(fname, parts[1])
        elif pi_type == "Person Name":
            _add_name_with_parts(fname, val)
        else:
            _add(fname, val)

    # scan all pair rows so no name is missed
    if redacting_names:
        df_all_pairs = pi_df[pi_df['PI Type'].isin(ALL_PAIR_TYPES)].dropna(subset=["File", "Detected Value"])
        for _, row in df_all_pairs.iterrows():
            fname = str(row["File"]).strip()
            val   = str(row["Detected Value"]).strip()
            name_part = val.split(", ", 1)[0].strip()
            _add_name_with_parts(fname, name_part)
        # also include standalone Person Name rows
        df_pn = pi_df[pi_df['PI Type'] == "Person Name"].dropna(subset=["File", "Detected Value"])
        for _, row in df_pn.iterrows():
            fname = str(row["File"]).strip()
            _add_name_with_parts(fname, str(row["Detected Value"]).strip())

    for fname in mapping:
        seen: set = set()
        deduped = []
        for v in mapping[fname]:
            if v not in seen:
                seen.add(v)
                deduped.append(v)
        mapping[fname] = deduped

    return mapping

def redact_excel_bytesio(excel_bytes: bytes, pi_values: List[str]) -> BytesIO:
    """Redact PII from an Excel file. Handles string, date (multiple formats), and numeric cells."""
    excel_stream = BytesIO(excel_bytes)
    excel_stream.seek(0)
    wb = openpyxl.load_workbook(excel_stream)

    pi_values_sorted = sorted(set(v for v in pi_values if v), key=len, reverse=True)

    # Pre-build alternate date formats so ISO dates also match other common formats
    date_alternates: Dict[str, List[str]] = {}
    for pi in pi_values_sorted:
        m = re.match(r'^(\d{4})-(\d{2})-(\d{2})$', pi)
        if m:
            y, mo, d = m.groups()
            date_alternates[pi] = [
                f"{mo}/{d}/{y}",    # MM/DD/YYYY (US Excel format)
                f"{d}/{mo}/{y}",    # DD/MM/YYYY (EU format)
                f"{y}/{mo}/{d}",    # YYYY/MM/DD
                f"{d}-{mo}-{y}",    # DD-MM-YYYY
                f"{mo}-{d}-{y}",    # MM-DD-YYYY
            ]

    def _cell_matches_pi(cell_str: str, pi: str) -> bool:
        # for short values (e.g. state codes), require a word boundary to avoid false matches
        if len(pi) <= 3 and pi.isalpha():
            import re as _re
            if not _re.search(r'(?<![A-Za-z0-9])' + _re.escape(pi) + r'(?![A-Za-z0-9])', cell_str, _re.IGNORECASE):
                return False
            return True
        if pi in cell_str:
            return True
        for alt in date_alternates.get(pi, []):
            if alt in cell_str:
                return True
        if pi.lower() in cell_str.lower():
            return True
        return False

    for sheet in wb.sheetnames:
        ws = wb[sheet]
        for row in ws.iter_rows():
            for cell in row:
                # string cells
                if isinstance(cell.value, str) and cell.value:
                    text = cell.value
                    changed = False
                    for pi in pi_values_sorted:
                        if _cell_matches_pi(text, pi):
                            # case-insensitive replace for all forms
                            import re as _re_cell
                            for alt in [pi] + date_alternates.get(pi, []):
                                new_text = _re_cell.sub(_re_cell.escape(alt), "[REDACTED]", text, flags=_re_cell.IGNORECASE)
                                if new_text != text:
                                    text = new_text
                                    changed = True
                    if changed:
                        cell.value = text

                # date/datetime cells
                elif hasattr(cell.value, 'strftime'):
                    cell_date = cell.value
                    for pi in pi_values_sorted:
                        # check ISO format
                        try:
                            iso = cell_date.strftime('%Y-%m-%d')
                            if iso == pi or pi in date_alternates and any(
                                a == cell_date.strftime('%m/%d/%Y') or
                                a == cell_date.strftime('%d/%m/%Y')
                                for a in date_alternates.get(pi, [])
                            ):
                                cell.value = "[REDACTED]"
                                break
                        except Exception:
                            pass

                # numeric cells (e.g. account numbers stored as numbers)
                elif cell.value is not None and not isinstance(cell.value, bool):
                    cell_str = str(cell.value)
                    for pi in pi_values_sorted:
                        if pi == cell_str or pi in cell_str:
                            cell.value = "[REDACTED]"
                            break

    out = BytesIO()
    wb.save(out)
    out.seek(0)
    return out

def safe_redact(text: str, value: str) -> str:
    """Case-insensitive regex redaction. More robust than plain str.replace() — handles spacing and formatting variants."""
    if not value or not text:
        return text
    pattern = re.escape(value)
    return re.sub(pattern, "[REDACTED]", text, flags=re.IGNORECASE)


def redact_docx_bytesio(docx_bytes: bytes, pi_values: List[str]) -> BytesIO:
    """Redact PII from a DOCX, handling ISO vs DD/MM/YYYY date format differences."""
    import re as _re
    stream = BytesIO(docx_bytes); stream.seek(0)
    doc = Document(stream)
    pi_sorted = sorted(set(v for v in pi_values if v), key=len, reverse=True)

    # Build alternate date forms so ISO dates match other common formats
    _date_alts: Dict[str, List[str]] = {}
    for pi in pi_sorted:
        m = _re.match(r"^(\d{4})-(\d{2})-(\d{2})$", pi)
        if m:
            y, mo, d = m.groups()
            _date_alts[pi] = [f"{d}/{mo}/{y}", f"{mo}/{d}/{y}", f"{d}-{mo}-{y}"]

    def _all_forms(pi):
        return [pi] + _date_alts.get(pi, [])

    def _redact_para(para):
        full = "".join(r.text for r in para.runs)
        replaced = full
        for pi in pi_sorted:
            for form in _all_forms(pi):
                escaped = _re.escape(form)
                # word boundaries for short alpha values (state codes, gender words) to prevent false matches
                if _re.fullmatch(r"[A-Za-z]{1,5}", form):
                    pattern = r"\b" + escaped + r"\b"
                else:
                    pattern = escaped
                if _re.search(pattern, replaced, _re.IGNORECASE):
                    replaced = _re.sub(pattern, "[REDACTED]", replaced, flags=_re.IGNORECASE)
        if replaced != full and para.runs:
            para.runs[0].text = replaced
            for r in para.runs[1:]: r.text = ""

    for para in doc.paragraphs:
        _redact_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _redact_para(para)
    out = BytesIO(); doc.save(out); out.seek(0)
    return out

try:
    import fitz as _fitz_module
except Exception:
    _fitz_module = None


def _get_local_fontsize(page, rect, default: float = 11.0) -> float:
    """Return the font size of the nearest non-redacted text span to rect."""
    best, dist = default, float("inf")
    for b in page.get_text("dict")["blocks"]:
        for line in b.get("lines", []):
            for span in line.get("spans", []):
                bb = span.get("bbox", [])
                if not bb:
                    continue
                centre_y = (bb[1] + bb[3]) / 2
                rect_cy  = (rect.y0 + rect.y1) / 2
                d = abs(centre_y - rect_cy)
                sz = span.get("size", 0)
                if d < dist and sz > 5 and "[REDACTED]" not in span.get("text", ""):
                    dist, best = d, sz
    return best


def redact_pdf_bytesio(pdf_bytes: bytes, pi_values: List[str]) -> BytesIO:
    """Redact PII from a PDF using PyMuPDF. Searches alternate date formats for full coverage."""
    try:
        import fitz
    except ImportError:
        print("[redact_pdf] PyMuPDF not installed — run: pip install pymupdf")
        out = BytesIO(pdf_bytes); out.seek(0); return out

    import re as _re
    # sort longer values first to prevent 'male' matching inside 'female'
    pi_sorted = sorted(set(v for v in pi_values if v), key=len, reverse=True)

    # build alternate date forms so ISO dates match other formats in the PDF
    all_search_terms = []
    for pi in pi_sorted:
        all_search_terms.append(pi)
        # add title-case and uppercase variants (PyMuPDF search is case-sensitive)
        if pi.islower() and len(pi) <= 15:
            all_search_terms.append(pi.title())
            all_search_terms.append(pi.upper())
        m = _re.match(r"^(\d{4})-(\d{2})-(\d{2})$", pi)
        if m:
            y, mo, d = m.groups()
            all_search_terms += [f"{d}/{mo}/{y}", f"{mo}/{d}/{y}", f"{d}-{mo}-{y}",
                                  f"{y}/{mo}/{d}", f"{y}-{mo}-{d}"]

    # normalize search terms
    normalized_terms = set()
    for term in all_search_terms:
        normalized_terms.add(term)
        collapsed = _re.sub(r'\s*-\s*', '-', term)
        if collapsed != term:
            normalized_terms.add(collapsed)

    # for addresses: also search city/state/ZIP portions separately
    address_extras = set()
    for term in list(normalized_terms):
        if ',' in term:
            parts = [p.strip() for p in term.split(',')]
            for part in parts[1:]:
                part = part.strip()
                if len(part) > 3:
                    address_extras.add(part)
                    pin = _re.search(r'\b\d{5,6}\b', part)
                    if pin:
                        address_extras.add(pin.group(0))
    normalized_terms.update(address_extras)

    # PDF tables sometimes break values across lines — add split fragments so we can find both parts
    split_fragments = set()
    for term in list(normalized_terms):
        # SSN split fragments
        m_ssn = _re.match(r'^(\d{3}-\d{2}-)(\d{4})$', term)
        if m_ssn:
            split_fragments.add(m_ssn.group(1))   # "253-75-"
            split_fragments.add(m_ssn.group(2))   # "9741"
        # ISO date split fragments
        m_iso = _re.match(r'^(\d{4}-)(\d{2}-\d{2})$', term)
        if m_iso:
            split_fragments.add(m_iso.group(1))   # "1966-"
            split_fragments.add(m_iso.group(2))   # "09-10"
        # date with slash — also add YYYY/ prefix
        m_yslash = _re.match(r'^(\d{4}/)(\d{2}/\d{2})$', term)
        if m_yslash:
            split_fragments.add(m_yslash.group(1))
            split_fragments.add(m_yslash.group(2))
        # phone split fragments
        m_phone = _re.match(r'^(\(\d{3}\))\s*(\d{3}-)(\d{4})$', term)
        if m_phone:
            split_fragments.add(m_phone.group(1))         # "(408)"
            split_fragments.add(m_phone.group(2))         # "321-"
            split_fragments.add(m_phone.group(3))         # "9136"
            split_fragments.add(m_phone.group(2) + m_phone.group(3))  # "321-9136"
        # Voter ID / MRN: skip — too ambiguous as fragments
    # only add fragments ≥6 chars to keep them specific
    normalized_terms.update(f for f in split_fragments if len(f) >= 6)

    all_search_terms = sorted(normalized_terms, key=len, reverse=True)

    # Separate full, fragment, and short word terms.
    # Short terms (≤5 chars like 'male', 'TX') are searched with word boundaries to avoid false matches.
    _FRAGMENT_MIN_LEN = 6
    full_terms  = [t for t in all_search_terms if len(t) >= 8]
    frag_terms  = [t for t in all_search_terms if len(t) < 8 and len(t) >= _FRAGMENT_MIN_LEN]
    # Short alpha-only terms that must match as whole words (gender values, state codes, etc.)
    short_word_terms = [t for t in all_search_terms if len(t) < _FRAGMENT_MIN_LEN and t.isalpha()]

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    for page in doc:
        # collect all match rects
        raw_rects: list = []
        # first pass: full-value terms
        found_full_terms: set = set()
        for term in full_terms:
            hits = page.search_for(term)
            if hits:
                found_full_terms.add(term)
                for rect in hits:
                    raw_rects.append(fitz.Rect(rect))
        # second pass: fragment terms — only when no full-value hit already covers it
        for term in frag_terms:
            # Skip if a longer term that contains this fragment was already found
            if any(term in ft for ft in found_full_terms):
                continue
            for rect in page.search_for(term):
                raw_rects.append(fitz.Rect(rect))
        # third pass: short whole-word terms matched via word boundaries (PyMuPDF has no built-in word-boundary search)
        if short_word_terms:
            _short_pattern = _re.compile(
                r"(?<![A-Za-z])(?:" + "|".join(_re.escape(t) for t in short_word_terms) + r")(?![A-Za-z])",
                _re.IGNORECASE,
            )
            _page_words = page.get_text("words")  # list of (x0,y0,x1,y1,word,block,line,word_idx)
            for _w in _page_words:
                _word_text = _w[4]
                if _short_pattern.fullmatch(_word_text.strip()):
                    raw_rects.append(fitz.Rect(_w[0], _w[1], _w[2], _w[3]))

        if not raw_rects:
            continue

        # sort top-to-bottom, left-to-right
        raw_rects.sort(key=lambda r: (round(r.y0, 1), r.x0))

        def _same_line(a, b, tol=4.0):
            return abs(a.y0 - b.y0) <= tol and abs(a.y1 - b.y1) <= tol

        def _adjacent(a, b, gap=6.0):
            return _same_line(a, b) and a.x0 <= b.x1 + gap and b.x0 <= a.x1 + gap

        # merge overlapping/adjacent rects on the same line
        merged: list = []
        for r in raw_rects:
            placed = False
            for i, m in enumerate(merged):
                if _adjacent(m, r):
                    merged[i] = fitz.Rect(
                        min(m.x0, r.x0), min(m.y0, r.y0),
                        max(m.x1, r.x1), max(m.y1, r.y1)
                    )
                    placed = True
                    break
            if not placed:
                merged.append(fitz.Rect(r))

        # draw each merged rect with a [REDACTED] label sized to the surrounding text
        for rect in merged:
            # sample nearby font size so [REDACTED] matches the document
            nearby_sz = 10.0  # safe default
            try:
                for blk in page.get_text("dict")["blocks"]:
                    for ln in blk.get("lines", []):
                        for sp in ln.get("spans", []):
                            bb = sp.get("bbox", [])
                            if bb and abs((bb[1]+bb[3])/2 - (rect.y0+rect.y1)/2) < 20:
                                sz = sp.get("size", 0)
                                if 6 < sz < 20:
                                    nearby_sz = sz
                                    break
            except Exception:
                pass

            # build a rect one text line high; wide enough to fit [REDACTED]
            line_h = nearby_sz * 1.25          # standard line height
            label = "[REDACTED]"
            # approximate char width (Helvetica ~0.52 × size)
            label_w = len(label) * nearby_sz * 0.52
            box = fitz.Rect(
                rect.x0,
                rect.y0,
                max(rect.x1, rect.x0 + label_w),
                rect.y0 + line_h,
            )
            page.add_redact_annot(
                box,
                text=label,
                fontsize=nearby_sz,
                text_color=(0, 0, 0),   # black text
                fill=(1, 1, 1),          # white background
                align=fitz.TEXT_ALIGN_CENTER,
            )
        page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)

    out = BytesIO()
    doc.save(out, garbage=4, deflate=True)
    doc.close()
    out.seek(0)
    return out

def process_source_files(source_files: List[Any], pi_dict: Dict[str, List[str]]) -> Dict[str, BytesIO]:
    """Redact PII from uploaded files (Excel, DOCX, PDF). Returns {file_name: BytesIO}."""
    redacted_output: Dict[str, BytesIO] = {}
    for f in source_files:
        fname = getattr(f, "name", "")
        ext = fname.lower().split(".")[-1]
        if ext not in {"xlsx", "xls", "docx", "pdf"}:
            continue
        try:
            f.seek(0)
        except Exception:
            pass
        original_bytes = f.read()
        pi_values = pi_dict.get(fname, [])
        try:
            if ext in {"xlsx", "xls"}:
                out_stream = redact_excel_bytesio(original_bytes, pi_values) if pi_values else BytesIO(original_bytes)
            elif ext == "docx":
                out_stream = redact_docx_bytesio(original_bytes, pi_values) if pi_values else BytesIO(original_bytes)
            elif ext == "pdf":
                out_stream = redact_pdf_bytesio(original_bytes, pi_values) if pi_values else BytesIO(original_bytes)
            else:
                out_stream = BytesIO(original_bytes)
        except Exception as e:
            print(f"✖ Redaction failed for {fname}: {e}")
            out_stream = BytesIO(original_bytes)
        out_stream.seek(0)
        redacted_output[fname] = out_stream
    return redacted_output

def to_excel(df: pd.DataFrame) -> bytes:
    output = BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name="PII Results")
    return output.getvalue()

# DOB parsing and validation
def parse_multiple_dates(text: Any) -> List[datetime]:
    if pd.isna(text):
        return []
    parts = [p.strip() for p in str(text).split(",")]
    parsed: List[datetime] = []
    for p in parts:
        dt = None
        # try ISO YYYY-MM-DD first (unambiguous)
        import re as _re_date
        if _re_date.match(r"\d{4}-\d{2}-\d{2}", p.strip()):
            try:
                dt = pd.to_datetime(p, dayfirst=False, errors="raise")
            except Exception:
                pass
        # then try DD/MM/YYYY
        if dt is None:
            try:
                dt = pd.to_datetime(p, dayfirst=True, errors="raise")
            except Exception:
                pass
        if dt is None:
            try:
                dt = pd.to_datetime(p, errors="raise")
            except Exception:
                pass
        if dt is None:
            for fmt in ["%d/%m/%Y", "%d-%m-%Y", "%d-%b-%Y",
                        "%m/%d/%Y", "%Y-%m-%d", "%d/%m/%y"]:
                try:
                    dt = datetime.strptime(p, fmt)
                    break
                except Exception:
                    continue
        if dt is not None:
            parsed.append(dt)
    return parsed

def is_probable_dob(date_obj: datetime) -> bool:
    today = datetime.today()
    age = today.year - date_obj.year - ((today.month, today.day) < (date_obj.month, date_obj.day))
    return 18 <= age <= 100

def filter_valid_dobs(value: Any) -> Any:
    if isinstance(value, list):
        return [d for d in value if is_probable_dob(d)]
    return value

# OCR for embedded images
def ocr_pil_image(img: Image.Image) -> str:
    """OCR a PIL image with preprocessing tuned for document scans. Returns extracted text."""
    # convert to grayscale
    gray = img.convert("L")

    # upscale small images for better accuracy (Tesseract works best at 300+ DPI)
    w, h = gray.size
    if w < 1500 or h < 1000:
        scale = max(2, 2400 // max(w, 1))
        gray = gray.resize((w * scale, h * scale), Image.LANCZOS)

    # slight sharpening for compressed/blurry scans
    gray = gray.filter(ImageFilter.SHARPEN)

    # run Tesseract with --psm 6 (uniform block of text)
    custom_config = r"--oem 1 --psm 6"
    try:
        text = pytesseract.image_to_string(gray, config=custom_config).strip()
    except Exception:
        # fallback: plain call
        text = pytesseract.image_to_string(gray).strip()
    return text

def extract_image_texts_from_pdf_pdfium(pdf_file, scale: float = 2.0) -> List[Dict[str, Any]]:
    records: List[Dict[str, Any]] = []
    doc_name = Path(getattr(pdf_file, "name", "document.pdf")).name

    # attempt 1: extract embedded image objects via pdfplumber
    try:
        try:
            pdf_file.seek(0)
        except Exception:
            pass
        with pdfplumber.open(pdf_file) as pdf_plumb:
            for page_idx, page in enumerate(pdf_plumb.pages, start=1):
                page_images = page.images or []
                img_idx = 0
                for img_info in page_images:
                    try:
                        # crop to image bounding box and convert to PIL
                        x0 = img_info.get("x0", 0)
                        top = img_info.get("top", 0)
                        x1 = img_info.get("x1", page.width)
                        bottom = img_info.get("bottom", page.height)
                        cropped = page.crop((x0, top, x1, bottom))
                        pil_img = cropped.to_image(resolution=200).original
                        text = ocr_pil_image(pil_img)
                        if text.strip():
                            img_idx += 1
                            records.append({
                                "doc_name": doc_name,
                                "page_or_sheet": f"PageImage {page_idx}",
                                "text": text,
                            })
                    except Exception:
                        continue
    except Exception:
        pass

    if records:
        return records

    # fallback: full-page render via pdfium (only when no image objects found)
    if pdfium is None:
        return records
    try:
        try:
            pdf_file.seek(0)
        except Exception:
            pass
        pdf = pdfium.PdfDocument(pdf_file)
    except Exception:
        return records

    n_pages = len(pdf)
    for i in range(n_pages):
        try:
            # skip OCR if the page already has extractable text
            _text_check = ""
            try:
                pdf_file.seek(0)
                from PyPDF2 import PdfReader as _PR
                _pr = _PR(pdf_file)
                if i < len(_pr.pages):
                    _text_check = _pr.pages[i].extract_text() or ""
            except Exception:
                pass
            # if the page has 100+ chars of extractable text, it's a text-layer page — skip OCR
            if len(_text_check.strip()) >= 100:
                continue

            page = pdf[i]
            bitmap = page.render(scale=scale, rotation=0)
            pil_image = bitmap.to_pil()
            if pil_image is None:
                continue
            text = ocr_pil_image(pil_image)
            records.append({"doc_name": doc_name, "page_or_sheet": f"PageImage {i + 1}", "text": text})
        except Exception:
            continue
    return records

def extract_image_texts_from_docx(docx_file) -> List[Dict[str, Any]]:
    records: List[Dict[str, Any]] = []
    doc_name = Path(getattr(docx_file, "name", "document.docx")).name

    try:
        docx_file.seek(0)
    except Exception:
        pass

    try:
        doc = Document(docx_file)
    except Exception:
        # fallback: save to a temp file and reopen
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            shutil.copyfileobj(docx_file, tmp)
            tmp_path = tmp.name
        try:
            doc = Document(tmp_path)
        finally:
            try:
                os.remove(tmp_path)
            except Exception:
                pass

    IMAGE_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
    idx = 0
    for rel in doc.part._rels.values():
        if rel.reltype == IMAGE_REL:
            idx += 1
            try:
                blob = rel.target_part.blob
                img = Image.open(BytesIO(blob))
                text = ocr_pil_image(img)
                records.append({"doc_name": doc_name, "page_or_sheet": f"InlineImage {idx}", "text": text})
            except Exception:
                pass

    return records

# OOXML namespaces for Excel XML parsing
NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "xdr": "http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
    "s": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
}
IMAGE_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
DRAWING_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/drawing"

def _read_xml(zf: ZipFile, path: str):
    try:
        with zf.open(path) as f:
            return ET.fromstring(f.read())
    except KeyError:
        return None

def _rels_path(part_path: str) -> str:
    base_dir = posixpath.dirname(part_path)
    base_name = posixpath.basename(part_path)
    return posixpath.join(base_dir, "_rels", base_name + ".rels")

def extract_image_texts_from_excel(xlsx_src) -> List[Dict[str, Any]]:
    """Extract embedded images from .xlsx/.xlsm and OCR them."""
    records: List[Dict[str, Any]] = []
    doc_name = Path(getattr(xlsx_src, "name", "workbook.xlsx")).name

    try:
        if hasattr(xlsx_src, "getvalue"):
            zf = ZipFile(BytesIO(xlsx_src.getvalue()))
        else:
            zf = ZipFile(xlsx_src)
    except Exception:
        return records

    with zf:
        wb_xml = _read_xml(zf, "xl/workbook.xml")
        wb_rels = _read_xml(zf, "xl/_rels/workbook.xml.rels")
        if wb_xml is None or wb_rels is None:
            return records

        rid_to_target = {}
        for rel in wb_rels.findall("rel:Relationship", NS):
            target = rel.get("Target", "")
            # absolute OOXML targets must be resolved from the zip root, not joined with 'xl/'
            if target.startswith("/"):
                resolved = posixpath.normpath(target.lstrip("/"))
            else:
                resolved = posixpath.normpath(posixpath.join("xl", target))
            rid_to_target[rel.get("Id")] = resolved

        sheet_part_to_name = {}
        for sheet in wb_xml.findall(".//s:sheet", NS):
            name = sheet.get("name")
            rid = sheet.get(f"{{{NS['r']}}}id")
            if not rid:
                continue
            target = rid_to_target.get(rid)
            if target:
                sheet_part_to_name[target] = name

        for sheet_part, sheet_name in sheet_part_to_name.items():
            sheet_rels_path = _rels_path(sheet_part)
            sh_rels = _read_xml(zf, sheet_rels_path)
            if sh_rels is None:
                continue

            drawing_targets = []
            for rel in sh_rels.findall("rel:Relationship", NS):
                if rel.get("Type") == DRAWING_REL:
                    target = rel.get("Target", "")
                    if target.startswith("/"):
                        drawing_path = posixpath.normpath(target.lstrip("/"))
                    else:
                        base_dir = posixpath.dirname(sheet_part)
                        drawing_path = posixpath.normpath(posixpath.join(base_dir, target))
                    drawing_targets.append(drawing_path)

            sheet_image_idx = 0
            for drawing_path in drawing_targets:
                drawing_xml = _read_xml(zf, drawing_path)
                if drawing_xml is None:
                    continue

                drawing_rels_path = _rels_path(drawing_path)
                drawing_rels = _read_xml(zf, drawing_rels_path)
                rid_to_img_target = {}
                if drawing_rels is not None:
                    for rel in drawing_rels.findall("rel:Relationship", NS):
                        if rel.get("Type") == IMAGE_REL:
                            rid_to_img_target[rel.get("Id")] = rel.get("Target")

                for blip in drawing_xml.findall(".//a:blip", NS):
                    embed_rid = blip.get(f"{{{NS['r']}}}embed")
                    if not embed_rid:
                        continue
                    img_target_rel = rid_to_img_target.get(embed_rid)
                    if not img_target_rel:
                        continue

                    if img_target_rel.startswith("/"):
                        media_path = img_target_rel.lstrip("/")
                    else:
                        draw_dir = posixpath.dirname(drawing_path)
                        media_path = posixpath.normpath(posixpath.join(draw_dir, img_target_rel)).lstrip("/")
                    try:
                        blob = zf.read(media_path)
                        img = Image.open(BytesIO(blob))
                        text = ocr_pil_image(img)
                        sheet_image_idx += 1
                        records.append({
                            "doc_name": doc_name,
                            "page_or_sheet": f"{sheet_name} : InlineImage {sheet_image_idx}",
                            "text": text,
                        })
                    except Exception:
                        pass

    return records

# Master OCR aggregator for image text
def extract_image_texts(file_paths: List[Any], pdf_scale: float = 2.0) -> pd.DataFrame:
    all_records: List[Dict[str, Any]] = []
    for f in file_paths:
        name = getattr(f, "name", "")
        ext = name.lower().split(".")[-1]
        try:
            if ext == "pdf":
                all_records.extend(extract_image_texts_from_pdf_pdfium(f, scale=pdf_scale))
            elif ext == "docx":
                all_records.extend(extract_image_texts_from_docx(f))
            elif ext in {"xlsx", "xls"}:
                all_records.extend(extract_image_texts_from_excel(f))
            else:
                pass  # skip unsupported
        except Exception as e:
            print(f"✖ Error processing in extract_image_texts for {name}: {e}")

    return pd.DataFrame(all_records, columns=["doc_name", "page_or_sheet", "text"])


# Detect PII from OCR output
def detect_pi_from_textframe(text_df: pd.DataFrame) -> pd.DataFrame:
    """Accepts a DataFrame with OCR output (doc_name, page_or_sheet, text) and returns PII detections in the standard format."""
    required_cols = {"doc_name", "page_or_sheet", "text"}
    if not required_cols.issubset(text_df.columns):
        raise ValueError(f"text_df is missing required columns: {required_cols - set(text_df.columns)}")

    compiled_us_dl = {state: re.compile(pattern, re.IGNORECASE) for state, pattern in US_DL_PATTERNS.items()}
    compiled_address = re.compile(ADDRESS_PATTERN, flags=re.IGNORECASE)
    compiled_company_id = ALNUM_COMPANY_ID_RE
    compiled_pi = {pi_type: re.compile(pattern) for pi_type, pattern in pi_PATTERNS.items()}
    use_spacy = get_nlp() is not None

    # file-level seen set for dedup
    seen: set[tuple] = set()

    def unique_preserve(seq: List[str]) -> List[str]:
        _s: set = set()
        return [s for s in seq if not (s in _s or _s.add(s))]

    records: List[Dict[str, Any]] = []

    for _, row in text_df.iterrows():
        file_name = row["doc_name"]
        page_or_sheet = row["page_or_sheet"]
        raw_text = row["text"] or ""
        cell_ref = "-"

        if not raw_text.strip():
            continue

        # OCR from image PDFs often splits mid-sentence
        ocr_text = raw_text

        # split into logical chunks for pair detection
        chunks = []
        for blk in re.split(r"(?:\n\s*\n)+|(?:^|\n)\s*#{1,6}\s.*?(?=\n|$)", ocr_text):
            b = (blk or "").strip()
            if not b:
                continue
            parts = [p.strip() for p in b.split("\n") if p.strip()]
            if len(parts) > 3:
                chunks.extend(parts)
            else:
                chunks.append(b)

        # Step 1: pair detection per chunk
        pair_emails:   set = set()
        pair_names:    set = set()
        pair_dobs:     set = set()
        pair_phones:   set = set()
        pair_addrs:    set = set()
        pair_cities:   set = set()
        pair_states:   set = set()
        pair_genders:  set = set()

        for i, blk in enumerate(chunks, 1):
            grouped = emit_grouped_records(file_name, f"{page_or_sheet}", cell_ref, blk, pair_only=True)
            for rec in grouped:
                _add_record(records, seen,
                    file_name=rec["File"],
                    page_sheet=rec["Page/Sheet"],
                    cell=rec["Cell"],
                    pi_type=rec["PI Type"],
                    value=rec["Detected Value"],
                    confidence=rec["Confidence (%)"],
                )
                parts = str(rec["Detected Value"]).split(", ", 1)
                name_part = parts[0].strip() if parts else ""
                val_part  = parts[1].strip() if len(parts) > 1 else ""
                pair_names.add(name_part.lower())
                if rec["PI Type"] == "Name, email":
                    pair_emails.add(val_part.lower())
                elif rec["PI Type"] == "Name, Date of Birth":
                    pair_dobs.add(val_part)
                elif rec["PI Type"] == "Name, Phone":
                    pair_phones.add(val_part)
                elif rec["PI Type"] == "Name, Address":
                    pair_addrs.add(val_part)
                elif rec["PI Type"] == "Name, City":
                    pair_cities.add(val_part.lower())
                elif rec["PI Type"] == "Name, State":
                    pair_states.add(val_part)
                elif rec["PI Type"] == "Name, Gender":
                    pair_genders.add(val_part.lower())

        # Step 2: all standalone PII on the full page text

        # SSN
        for v in find_ssn_in_text(ocr_text):
            _add_record(records, seen, file_name=file_name, page_sheet=page_or_sheet,
                        cell=cell_ref, pi_type="SSN Number", value=v, confidence=100)

        # person names (suppress if already in a pair)
        max_for_names = 10 if str(page_or_sheet).lower().startswith("pageimage") else 20
        for name in extract_person_names(ocr_text, max_names=max_for_names):
            if name.lower() not in pair_names:
                _add_record(records, seen, file_name=file_name, page_sheet=page_or_sheet,
                            cell=cell_ref, pi_type="Person Name", value=name,
                            confidence=95 if use_spacy else 85)

        # phone (suppress if already in a pair)
        for m in re.finditer(pi_PATTERNS["Phone"], ocr_text):
            v = m.group(0)
            s = m.start()
            if s > 0 and ocr_text[s-1] == "(" and not v.startswith("("):
                v = "(" + v
            if re.match(r"^\d{3}-\d{2}-\d{4}$", v.strip()):
                continue
            if '.' in v and re.search(r'\d+\.\d+\.\d+', v):
                continue
            if v in pair_phones:
                continue
            _add_record(records, seen, file_name=file_name, page_sheet=page_or_sheet,
                        cell=cell_ref, pi_type="Phone", value=v, confidence=88)

        # email (suppress if already in a pair)
        for m in re.finditer(pi_PATTERNS["email"], ocr_text):
            v = m.group(0)
            if v.lower() not in pair_emails:
                _add_record(records, seen, file_name=file_name, page_sheet=page_or_sheet,
                            cell=cell_ref, pi_type="email", value=v, confidence=100)

        # date of birth (suppress if already in a pair; filter non-plausible dates)
        dob_candidates: List[str] = []
        for cp in DOB_PATTERNS:
            dob_candidates.extend(m.group(0) for m in cp.finditer(ocr_text))
        dob_unique = unique_preserve(dob_candidates)
        for dv in dob_unique:
            if dv in pair_dobs:
                continue
            # only emit if plausible as a birth date
            try:
                import dateutil.parser as _dup
                _parsed = _dup.parse(dv, dayfirst=False)
                if not is_probable_dob(_parsed):
                    continue
            except Exception:
                pass  # if we can't parse it, allow it through
            _add_record(records, seen, file_name=file_name, page_sheet=page_or_sheet,
                        cell=cell_ref, pi_type="Date of Birth", value=dv, confidence=90)

        # credit/debit card (Luhn-validated)
        for m in re.finditer(pi_PATTERNS["Credit/Debit Card"], ocr_text):
            v = m.group(0)
            if passes_luhn(v):
                _add_record(records, seen, file_name=file_name, page_sheet=page_or_sheet,
                            cell=cell_ref, pi_type="Credit/Debit Card", value=v, confidence=92)

        # IP Address
        for m in re.finditer(pi_PATTERNS["IP Address"], ocr_text):
            _add_record(records, seen, file_name=file_name, page_sheet=page_or_sheet,
                        cell=cell_ref, pi_type="IP Address", value=m.group(0), confidence=95)

        # US Passport — requires a passport keyword nearby
        for m in re.finditer(pi_PATTERNS["US Passport Number"], ocr_text):
            v = m.group(0).strip()
            if re.match(r"^[A-Z]\d{8}$", v) or PASSPORT_9D_RE.match(v):
                if PASSPORT_KEYWORD.search(ocr_text):
                    _add_record(records, seen, file_name=file_name, page_sheet=page_or_sheet,
                                cell=cell_ref, pi_type="US Passport Number", value=v, confidence=95)

        # driving license (context-gated)
        for state, full_value, _ in iter_dl_matches(ocr_text, col_name=None, require_context=True):
            _add_record(records, seen, file_name=file_name, page_sheet=page_or_sheet,
                        cell=cell_ref, pi_type=f"Driving License ({state})", value=full_value, confidence=95)

        # Address
        flat = re.sub(r"[\r\n\t]+", " ", ocr_text)
        for addr in unique_preserve([m.strip() for m in compiled_address.findall(flat) if m and m.strip()]):
            _add_record(records, seen, file_name=file_name, page_sheet=page_or_sheet,
                        cell=cell_ref, pi_type="Address", value=addr, confidence=80)
        # City+State+ZIP fragment
        _CITY_ST_ZIP = re.compile(r"\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?),\s*([A-Z]{2})\s+(\d{5,6})\b")
        for m in _CITY_ST_ZIP.finditer(flat):
            _add_record(records, seen, file_name=file_name, page_sheet=page_or_sheet,
                        cell=cell_ref, pi_type="Address", value=m.group(0).strip(), confidence=78)

        # State
        states_detected = detect_states(ocr_text)
        if states_detected:
            _add_record(records, seen, file_name=file_name, page_sheet=page_or_sheet,
                        cell=cell_ref, pi_type="State", value=states_detected, confidence=75)

        # Account Number (requires banking keyword)
        _indian_phone_re = re.compile(r"^[6-9]\d{9}$")
        for m in re.finditer(pi_PATTERNS["Account Number"], ocr_text):
            v = m.group(0)
            digits = re.sub(r"\D", "", v)
            if passes_luhn(v) or _indian_phone_re.match(digits):
                continue
            if _has_account_context(ocr_text):
                if _is_routing_number(v, ocr_text):
                    _add_record(records, seen, file_name=file_name, page_sheet=page_or_sheet,
                                cell=cell_ref, pi_type="Routing Number", value=v, confidence=90)
                else:
                    _add_record(records, seen, file_name=file_name, page_sheet=page_or_sheet,
                                cell=cell_ref, pi_type="Account Number", value=v,
                                confidence=90 if len(digits) >= 12 else 80)

        # MRN — medical keyword required
        for m in MRN_RE.finditer(ocr_text):
            if _has_mrn_context(ocr_text):
                _add_record(records, seen, file_name=file_name, page_sheet=page_or_sheet,
                            cell=cell_ref, pi_type="MRN", value=m.group(1) or m.group(2) or "", confidence=90)

        # Insurance ID
        for m in INSURANCE_ID_RE.finditer(ocr_text):
            _add_record(records, seen, file_name=file_name, page_sheet=page_or_sheet,
                        cell=cell_ref, pi_type="Insurance ID", value=m.group(1) or m.group(2) or "", confidence=88)

        # Aadhaar
        for m in re.finditer(pi_PATTERNS["Aadhaar Number"], ocr_text):
            _add_record(records, seen, file_name=file_name, page_sheet=page_or_sheet,
                        cell=cell_ref, pi_type="Aadhaar Number", value=m.group(0), confidence=97)

        # PAN Card
        for m in re.finditer(pi_PATTERNS["PAN Card"], ocr_text):
            _add_record(records, seen, file_name=file_name, page_sheet=page_or_sheet,
                        cell=cell_ref, pi_type="PAN Card", value=m.group(0), confidence=98)

        # Company ID — skip IDs already captured inside a DL value
        detected_dl_values = {r["Detected Value"] for r in records
                              if r.get("PI Type","").startswith("Driving License") and r.get("File")==file_name}
        for cid in compiled_company_id.findall(ocr_text):
            if any(cid in dl_val for dl_val in detected_dl_values):
                continue  # already part of a DL value
            _add_record(records, seen, file_name=file_name, page_sheet=page_or_sheet,
                        cell=cell_ref, pi_type="Company ID", value=cid, confidence=90)

        # Gender
        for val in detect_gender_values(ocr_text):
            _add_record(records, seen, file_name=file_name, page_sheet=page_or_sheet,
                        cell=cell_ref, pi_type="Gender", value=val, confidence=85)

    return pd.DataFrame(records, columns=["File", "Page/Sheet", "Cell", "PI Type", "Detected Value", "Confidence (%)"])

def clean_pi_data(pi_df: pd.DataFrame) -> pd.DataFrame:
    """Deduplicate and standardize results. Drops rows below 40% confidence (weak passports, account numbers without context, etc.)."""
    if pi_df is None or pi_df.empty:
        return pi_df

    # normalize column name capitalization
    cols = {c: c.strip() for c in pi_df.columns}
    pi_df = pi_df.rename(columns=cols)
    if "PI Type" not in pi_df.columns and "pi Type" in pi_df.columns:
        pi_df = pi_df.rename(columns={"pi Type": "PI Type"})

    df = pi_df.dropna(subset=["Detected Value"]).copy()
    df.drop_duplicates(inplace=True)

    # drop low-confidence rows before grouping
    if "Confidence (%)" in df.columns:
        df = df[df["Confidence (%)"] >= 40].copy()
    # remove internal helper types
    df = df[~df["PI Type"].isin(["City_SKIP"])].copy() if "PI Type" in df.columns else df

    if df.empty:
        return df

    if "PI Type" in df.columns:
        pair_email_keys: set = set()
        for _, r in df[df["PI Type"] == "Name, email"].iterrows():
            parts = str(r["Detected Value"]).split(", ", 1)
            if len(parts) == 2:
                pair_email_keys.add((str(r["File"]), parts[1].strip()))
        mask_dup_email = df.apply(
            lambda r: r["PI Type"] == "email" and
                      (str(r["File"]), str(r["Detected Value"]).strip()) in pair_email_keys,
            axis=1
        )
        df = df[~mask_dup_email].copy()

    grouped = (
        df.groupby(["File", "Page/Sheet", "PI Type", "Detected Value"], dropna=False)
          .agg(**{"Confidence (%)": ("Confidence (%)", "max"), "Occurrence": ("File", "count")})
          .reset_index()
    )

    # DOB cleanup
    dob_mask = grouped["PI Type"].eq("Date of Birth")
    if dob_mask.any():
        parsed = grouped.loc[dob_mask, "Detected Value"].apply(parse_multiple_dates)
        valid = parsed.apply(filter_valid_dobs)
        # drop rows with no valid DOB remaining
        keep_mask = valid.apply(lambda x: isinstance(x, list) and len(x) > 0)
        grouped = pd.concat(
            [grouped.loc[~dob_mask], grouped.loc[dob_mask].loc[keep_mask].assign(**{
                "Detected Value": valid[keep_mask].apply(lambda xs: ", ".join(sorted({d.strftime('%Y-%m-%d') for d in xs})))
            })],
            ignore_index=True
        )

    # drop empty detected values
    grouped = grouped[grouped["Detected Value"].astype(str).str.strip() != ""]

    # drop single-word Person Name values (last-name-only false positives)
    if "Person Name" in grouped["PI Type"].values:
        grouped = grouped[~(
            (grouped["PI Type"] == "Person Name") &
            grouped["Detected Value"].apply(
                lambda v: len(str(v).split()) == 1
            )
        )].copy()

    # Address: drop substrings of longer address values
    if "Address" in grouped["PI Type"].values:
        addr_vals = grouped[grouped["PI Type"] == "Address"]["Detected Value"].tolist()
        addr_vals_sorted = sorted(addr_vals, key=len, reverse=True)  # longest first
        keep = []
        for val in addr_vals_sorted:
            # keep if not a substring of an already-kept longer address
            if not any(str(val) in str(longer) and str(val) != str(longer) for longer in keep):
                keep.append(val)
        keep_set = set(keep)
        grouped = grouped[
            (grouped["PI Type"] != "Address") | (grouped["Detected Value"].isin(keep_set))
        ].copy()


    def _extract_para_num(val: str) -> int:
        m = re.search(r'(\d+)', str(val))
        return int(m.group(1)) if m else 0

    grouped["_para_num"] = grouped["Page/Sheet"].apply(_extract_para_num)
    grouped = grouped.sort_values(by=["PI Type", "_para_num", "Detected Value"]).drop(columns=["_para_num"])
    grouped = grouped.reset_index(drop=True)

    # if the same 10-digit number appears as both Indian Phone and Account Number, prefer Account Number
    if "Indian Phone" in grouped["PI Type"].values and "Account Number" in grouped["PI Type"].values:
        acct_digits = set(
            re.sub(r"\D", "", str(v))
            for v in grouped[grouped["PI Type"] == "Account Number"]["Detected Value"]
        )
        grouped = grouped[~(
            (grouped["PI Type"] == "Indian Phone") &
            grouped["Detected Value"].apply(lambda v: re.sub(r"\D", "", str(v)) in acct_digits)
        )].copy()

    # remove duplicate (the block above handles this already)
    if "Indian Phone" in grouped["PI Type"].values and "Account Number" in grouped["PI Type"].values:
        acct_digits = set(
            re.sub(r"\D", "", str(v))
            for v in grouped[grouped["PI Type"] == "Account Number"]["Detected Value"]
        )
        grouped = grouped[~(
            (grouped["PI Type"] == "Indian Phone") &
            grouped["Detected Value"].apply(lambda v: re.sub(r"\D", "", str(v)) in acct_digits)
        )].copy()

    # if the same number appears as both Indian Phone and Phone, keep only Indian Phone
    if "Indian Phone" in grouped["PI Type"].values and "Phone" in grouped["PI Type"].values:
        import re as _re
        indian_last10 = set(
            _re.sub(r"\D", "", str(v))[-10:]
            for v in grouped[grouped["PI Type"] == "Indian Phone"]["Detected Value"]
            if len(_re.sub(r"\D", "", str(v))) >= 10
        )
        def _dup_phone(row):
            if row["PI Type"] != "Phone":
                return False
            d = _re.sub(r"\D", "", str(row["Detected Value"]))
            return len(d) >= 10 and d[-10:] in indian_last10
        grouped = grouped[~grouped.apply(_dup_phone, axis=1)].copy()


    if "PI Type" in grouped.columns:
        # drop standalone email if it already appears in a Name+email pair
        pair_email_keys: set = set()
        for _, r in grouped[grouped["PI Type"] == "Name, email"].iterrows():
            parts = str(r["Detected Value"]).split(", ", 1)
            if len(parts) == 2:
                pair_email_keys.add((str(r["File"]), parts[1].strip()))
        mask_dup = grouped.apply(
            lambda r: r["PI Type"] == "email" and
                      (str(r["File"]), str(r["Detected Value"]).strip()) in pair_email_keys,
            axis=1
        )
        grouped = grouped[~mask_dup].copy()

    if "PI Type" in grouped.columns and "Person Name" in grouped["PI Type"].values:
        _ALL_PAIR_TYPES = {"Name, email", "Name, Date of Birth", "Name, Company ID",
                           "Name, Phone", "Name, Address", "Name, City", "Name, State", "Name, Gender"}
        _pair_name_keys: set = set()
        _pair_name_words: set = set()  # individual words from paired full names
        for _, _r in grouped[grouped["PI Type"].isin(_ALL_PAIR_TYPES)].iterrows():
            _name_part = str(_r["Detected Value"]).split(", ", 1)[0].strip()
            _pair_name_keys.add(str(_r["File"]) + "|" + _name_part.lower())
            # add each word of the full name
            for _w in _name_part.split():
                _pair_name_words.add(str(_r["File"]) + "|" + _w.lower())

        def _standalone_name_covered(r) -> bool:
            if r["PI Type"] != "Person Name":
                return False
            val = str(r["Detected Value"]).strip()
            file_key = str(r["File"])
            # exact match (case-insensitive)
            if file_key + "|" + val.lower() in _pair_name_keys:
                return True
            # Single-word first/last name that is a word in any paired full name
            if len(val.split()) == 1 and file_key + "|" + val.lower() in _pair_name_words:
                return True
            return False

        grouped = grouped[~grouped.apply(_standalone_name_covered, axis=1)].copy()

    if "PI Type" in grouped.columns and "Date of Birth" in grouped["PI Type"].values:
        _pair_dob_keys: set = set()
        for _, _r in grouped[grouped["PI Type"] == "Name, Date of Birth"].iterrows():
            _dob_part = str(_r["Detected Value"]).split(", ", 1)
            if len(_dob_part) == 2:
                _pair_dob_keys.add((str(_r["File"]), _dob_part[1].strip()))
        grouped = grouped[~(
            (grouped["PI Type"] == "Date of Birth") &
            grouped.apply(
                lambda r: (str(r["File"]), str(r["Detected Value"]).strip()) in _pair_dob_keys,
                axis=1,
            )
        )].copy()

    # suppress standalone Phone if already in a Name,Phone pair
    if "PI Type" in grouped.columns:
        _pair_phone_keys: set = set()
        for _, _r in grouped[grouped["PI Type"] == "Name, Phone"].iterrows():
            _p = str(_r["Detected Value"]).split(", ", 1)
            if len(_p) == 2:
                _pair_phone_keys.add((str(_r["File"]), _p[1].strip()))
        for _pt in ("Phone", "Indian Phone"):
            grouped = grouped[~(
                (grouped["PI Type"] == _pt) &
                grouped.apply(
                    lambda r: (str(r["File"]), str(r["Detected Value"]).strip()) in _pair_phone_keys,
                    axis=1,
                )
            )].copy()

    # suppress standalone Address if already in a Name,Address pair, including city/state/ZIP fragments
    if "PI Type" in grouped.columns:
        _pair_addr_vals: set = set()
        _pair_addr_files: set = set()
        for _, _r in grouped[grouped["PI Type"] == "Name, Address"].iterrows():
            _a = str(_r["Detected Value"]).split(", ", 1)
            if len(_a) == 2:
                _pair_addr_vals.add((str(_r["File"]), _a[1].strip()))
                _pair_addr_files.add(str(_r["File"]))
        # pattern for city/state/zip fragments without a street number
        _city_state_frag_re = re.compile(
            r"^[A-Za-z][A-Za-z\s]+,\s*[A-Z]{2}\s+\d{5,6}$|"  # City, ST ZIPCODE
            r"^[A-Za-z][A-Za-z\s]+\s+\d{5,6}$|"               # City ZIPCODE (Indian style)
            r"^(?:Floor|Suite|Apt|Unit|Flat)\s+\S+"             # Floor 3 / Suite 755 alone
        )
        def _addr_covered(r):
            if r["PI Type"] != "Address":
                return False
            rv = str(r["Detected Value"]).strip()
            fk = str(r["File"])
            # exact match in pair addresses
            if (fk, rv) in _pair_addr_vals:
                return True
            # also suppress if it's a sub-fragment of a paired address
            if any(fk == pk and rv in pv for pk, pv in _pair_addr_vals):
                return True
            # city/state/zip fragments for files that already have a Name,Address pair
            if fk in _pair_addr_files and _city_state_frag_re.match(rv):
                return True
            return False
        grouped = grouped[~grouped.apply(_addr_covered, axis=1)].copy()

    # suppress standalone City if already in a Name,City pair
    if "PI Type" in grouped.columns and "City" in grouped["PI Type"].values:
        _pair_city_keys: set = set()
        for _, _r in grouped[grouped["PI Type"] == "Name, City"].iterrows():
            _c = str(_r["Detected Value"]).split(", ", 1)
            if len(_c) == 2:
                _pair_city_keys.add((str(_r["File"]), _c[1].strip().lower()))
        grouped = grouped[~(
            (grouped["PI Type"] == "City") &
            grouped.apply(
                lambda r: (str(r["File"]), str(r["Detected Value"]).strip().lower()) in _pair_city_keys,
                axis=1,
            )
        )].copy()

    # suppress standalone State if already in a Name,State pair or inside a Name,Address value
    if "PI Type" in grouped.columns and "State" in grouped["PI Type"].values:
        _pair_state_keys: set = set()
        for _, _r in grouped[grouped["PI Type"] == "Name, State"].iterrows():
            _s = str(_r["Detected Value"]).split(", ", 1)
            if len(_s) == 2:
                _pair_state_keys.add((str(_r["File"]), _s[1].strip()))
        # also collect state codes from Name,Address values
        _addr_state_codes: set = set()
        for _, _r in grouped[grouped["PI Type"] == "Name, Address"].iterrows():
            _addr_val = str(_r["Detected Value"])
            _fk = str(_r["File"])
            import re as _re2
            for _code in _re2.findall(r'\b([A-Z]{2})\b', _addr_val):
                _addr_state_codes.add((_fk, _code))
        def _state_covered(r):
            if r["PI Type"] != "State":
                return False
            fk = str(r["File"])
            sv = str(r["Detected Value"]).strip()
            if (fk, sv) in _pair_state_keys:
                return True
            # handle comma-joined state codes like 'TN, CA'
            for code in sv.replace(",", " ").split():
                if (fk, code.strip()) in _addr_state_codes:
                    return True
            return False
        grouped = grouped[~grouped.apply(_state_covered, axis=1)].copy()

    # suppress standalone Company ID if already in a Name,Company ID pair
    if "PI Type" in grouped.columns and "Company ID" in grouped["PI Type"].values:
        _pair_cid_keys: set = set()
        for _, _r in grouped[grouped["PI Type"] == "Name, Company ID"].iterrows():
            _c = str(_r["Detected Value"]).split(", ", 1)
            if len(_c) == 2:
                _pair_cid_keys.add((str(_r["File"]), _c[1].strip()))
        grouped = grouped[~(
            (grouped["PI Type"] == "Company ID") &
            grouped.apply(
                lambda r: (str(r["File"]), str(r["Detected Value"]).strip()) in _pair_cid_keys,
                axis=1,
            )
        )].copy()

    # drop Phone values that look like IP address fragments
    import re as _re_phone
    if "Phone" in grouped["PI Type"].values:
        grouped = grouped[~(
            (grouped["PI Type"] == "Phone") &
            grouped["Detected Value"].apply(
                lambda v: bool(_re_phone.match(r"^\d+\.\d+\.\d+", str(v)))
            )
        )].copy()

    if "Indian Phone" in grouped["PI Type"].values and "Phone" in grouped["PI Type"].values:
        _ind_last10 = set(
            re.sub(r"\D","",str(v))[-10:]
            for v in grouped[grouped["PI Type"]=="Indian Phone"]["Detected Value"]
            if len(re.sub(r"\D","",str(v))) >= 10
        )
        def _is_dup_phone(row):
            if row["PI Type"] != "Phone": return False
            d = re.sub(r"\D","",str(row["Detected Value"]))
            return len(d) >= 10 and d[-10:] in _ind_last10
        grouped = grouped[~grouped.apply(_is_dup_phone, axis=1)].copy()

    # drop Indian Phone if same number already flagged as Account Number
    if "Indian Phone" in grouped["PI Type"].values and "Account Number" in grouped["PI Type"].values:
        _acct_digits = set(re.sub(r"\D","",str(v)) for v in grouped[grouped["PI Type"]=="Account Number"]["Detected Value"])
        grouped = grouped[~(
            (grouped["PI Type"]=="Indian Phone") &
            grouped["Detected Value"].apply(lambda v: re.sub(r"\D","",str(v)) in _acct_digits)
        )].copy()

    import re as _re2
    def _pg_num(s):
        nums = _re2.findall(r"\d+", str(s))
        return int(nums[0]) if nums else 0
    try:
        grouped = grouped.copy()
        grouped["_pg_sort"] = grouped["Page/Sheet"].apply(_pg_num)
        grouped = grouped.sort_values(
            ["PI Type", "_pg_sort", "Detected Value"]
        ).drop(columns=["_pg_sort"]).reset_index(drop=True)
    except Exception:
        pass

    return grouped
